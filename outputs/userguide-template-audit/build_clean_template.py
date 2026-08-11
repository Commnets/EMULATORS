from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_field(paragraph, instruction: str, placeholder: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        paragraph.add_run()._r.append(node)


def page_break(document: Document) -> None:
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build(source: Path, destination: Path) -> None:
    document = Document(source)
    body = document._element.body
    section_properties = body.sectPr
    children = list(body)
    cover = next((child for child in children if child.tag == qn("w:sdt")), None)
    if cover is None:
        raise RuntimeError("The source document does not contain the expected cover-page content control")

    for child in children:
        if child is not cover and child is not section_properties:
            body.remove(child)

    heading = document.add_paragraph("Índice", style="Heading 1")
    heading.paragraph_format.keep_with_next = True
    toc = document.add_paragraph()
    add_field(
        toc,
        ' TOC \\h \\z \\t "Título 1;2;Título 2;3;Título 3;4;Título;1" ',
        "Actualice los campos para generar el índice.",
    )

    page_break(document)
    document.add_paragraph("Prólogo", style="Title")
    document.add_paragraph("[Escriba aquí el prólogo.]")

    page_break(document)
    document.add_paragraph("[Título del bloque de contenido]", style="Title")
    document.add_paragraph("[Desarrolle aquí el contenido. Duplique este bloque tantas veces como sea necesario.]")

    page_break(document)
    document.add_paragraph("Anexos", style="Title")
    document.add_paragraph("[Añada aquí los anexos o elimine este texto si no son necesarios.]")

    page_break(document)
    document.add_paragraph("Ilustraciones", style="Heading 1")
    figures = document.add_paragraph(style="table of figures")
    add_field(figures, ' TOC \\h \\z \\c "Ilustración" ', "Actualice los campos para generar la lista de ilustraciones.")

    page_break(document)
    document.add_paragraph("Ecuaciones", style="Heading 1")
    equations = document.add_paragraph(style="table of figures")
    add_field(equations, ' TOC \\h \\z \\c "Ecuación" ', "Actualice los campos para generar la lista de ecuaciones.")

    document.core_properties.title = "Título del documento"
    document.core_properties.subject = "Subtítulo"
    document.core_properties.author = "Ignacio Cea Forniés"
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_clean_template.py SOURCE.docx DESTINATION.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))
