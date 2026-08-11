from pathlib import Path

from docx import Document


document = Document(r"C:\Workspaces\EMULATORS\docs\C64Emulator_UserGuide.docx")
lines = []
for index, paragraph in enumerate(document.paragraphs):
    style = paragraph.style.name if paragraph.style is not None else ""
    text = paragraph.text.strip()
    if style in {"Title", "Heading 1", "Heading 2", "Heading 3"}:
        lines.append(f"{index}\t{style}\t{text}")

lines.append("\nW BLOCKS")
for index, paragraph in enumerate(document.paragraphs):
    if paragraph.text.strip() == "/w":
        for item in document.paragraphs[index:index + 12]:
            lines.append(f"{item.style.name}\t{item.text}")
        lines.append("---")

Path(r"C:\Workspaces\EMULATORS\outputs\vic20-user-guide-work\guide-structure.txt").write_text(
    "\n".join(lines), encoding="utf-8"
)
