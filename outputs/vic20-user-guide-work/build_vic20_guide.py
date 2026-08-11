from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Pt


GLOBAL_REPLACEMENTS = (
    ("C64::C64Emulator", "VIC20::Emulator"),
    ("C64EmulatorC", "VIC20EmulatorC"),
    ("C64Emulator", "VIC20Emulator"),
    ("C64Emulator_UserGuide.docx", "VIC20Emulator_UserGuide.docx"),
    ("C64EmulatorC.exe", "VIC20EmulatorC.exe"),
    ("C64Emulator.exe", "VIC20Emulator.exe"),
    ("C64::", "VIC20::"),
    ("C64-Specific", "VIC20-Specific"),
    ("C64-specific", "VIC20-specific"),
    ("C64 builder", "VIC20 builder"),
    ("Commodore 64", "Commodore VIC-20"),
    ("6510", "6502"),
    ("c64.log", "vic20.log"),
    ("c64-screen.png", "vic20-screen.png"),
    ("C64", "VIC-20"),
)


VIC20_COMMANDS = (
    {
        "name": "VIA1",
        "aliases": "CVIA1.",
        "meaning": "Shows the complete state of VIA #1.",
        "syntax": "VIA1",
        "parameters": ["No parameters."],
        "examples": ["VIA1"],
        "result": "Returns the VIC-20 VIA #1 register, timer, port and interrupt state.",
    },
    {
        "name": "VIA2",
        "aliases": "CVIA2.",
        "meaning": "Shows the complete state of VIA #2.",
        "syntax": "VIA2",
        "parameters": ["No parameters."],
        "examples": ["VIA2"],
        "result": "Returns the VIC-20 VIA #2 register, timer, port and interrupt state.",
    },
    {
        "name": "SCREENDUMP",
        "aliases": "CSCREENDUMP.",
        "meaning": "Returns the current text-screen memory in hexadecimal.",
        "syntax": "SCREENDUMP",
        "parameters": ["No parameters."],
        "examples": ["SCREENDUMP"],
        "result": "Returns the VIC-I screen-memory snapshot visible to the active CPU.",
    },
    {
        "name": "COLORDUMP",
        "aliases": "CCOLORDUMP.",
        "meaning": "Returns the current color memory in hexadecimal.",
        "syntax": "COLORDUMP",
        "parameters": ["No parameters."],
        "examples": ["COLORDUMP"],
        "result": "Returns the VIC-I color-memory snapshot visible to the active CPU.",
    },
    {
        "name": "CHARSDRAW",
        "aliases": "CCHARSDRAW.",
        "meaning": "Renders textual drawings of all or selected character glyphs.",
        "syntax": "CHARSDRAW [0..255 ...]",
        "parameters": [
            "Character codes: optional values from 0 through 255; duplicates are coalesced. With none, all characters are rendered."
        ],
        "examples": ["CHARSDRAW", "CHARSDRAW 1 2 65"],
        "result": "Returns textual drawings generated from VIC-I character data.",
    },
    {
        "name": "LIGHTPEN",
        "aliases": "CLIGHTPEN.",
        "meaning": "Enables or disables mouse-driven light-pen emulation.",
        "syntax": "LIGHTPEN ON|OFF",
        "parameters": ["ON: enable.", "OFF: disable."],
        "examples": ["LIGHTPEN ON", "LIGHTPEN OFF"],
        "result": "The mouse moves the VIC-I light pen and its left button presses it. This implementation can coexist with joysticks.",
    },
)


INHERITED_RESULTS = {
    "VICII": (
        "Requests VIC-II state.",
        "Accepted by the inherited builder, but the VIC-20 has no VIC-II, so the command returns no chip data.",
    ),
    "VICIIEVENTS": (
        "Enables or disables VIC-II event visualization.",
        "Accepted by the inherited builder, but the VIC-20 has no VIC-II, so the command has no effect.",
    ),
    "SID": (
        "Requests SID state.",
        "Accepted by the inherited builder, but the VIC-20 has no separate SID chip; VIC-I provides its sound generation, so no SID data is returned.",
    ),
    "SIDW": (
        "Changes the active SID sound wrapper.",
        "Accepted by the inherited builder, but the VIC-20 has no SID chip, so the sound implementation is unchanged.",
    ),
    "VICI": (
        "Shows VIC-I video, raster and sound state.",
        "Returns the VIC-I state used by the VIC-20.",
    ),
    "TED": (
        "Requests TED state.",
        "Accepted by the inherited builder, but the VIC-20 has no TED chip, so the command returns no chip data.",
    ),
    "TEDEVENTS": (
        "Enables or disables TED event visualization.",
        "Accepted by the inherited builder, but the VIC-20 has no TED chip, so the command has no effect.",
    ),
    "VIA": (
        "Requests the generic Commodore VIA state.",
        "The VIC-20 registers its controllers as VIA1 and VIA2 rather than the generic VIA identifier; use VIA1 or VIA2 for deterministic output.",
    ),
    "CIA": (
        "Requests generic Commodore CIA state.",
        "Accepted by the inherited builder, but the VIC-20 uses VIA1 and VIA2 and has no CIA, so no CIA data is returned.",
    ),
}


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_before(anchor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addprevious(element)
    paragraph = Paragraph(element, anchor._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def insert_labeled(anchor: Paragraph, label: str, text: str) -> Paragraph:
    paragraph = insert_before(anchor, style="Normal")
    paragraph.add_run(label).bold = True
    paragraph.add_run(text)
    return paragraph


def add_command_block(anchor: Paragraph, command: dict[str, object]) -> None:
    heading = insert_before(anchor, str(command["name"]), "Heading 2")
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.keep_with_next = True
    insert_labeled(anchor, "Availability: ", "Local console and remote /o channel.")
    insert_labeled(anchor, "Accepted aliases: ", str(command["aliases"]))
    insert_labeled(anchor, "Meaning: ", str(command["meaning"]))
    insert_labeled(anchor, "General syntax: ", "")
    insert_before(anchor, str(command["syntax"]), "Code")
    insert_labeled(anchor, "Parameters: ", "")
    for parameter in command["parameters"]:
        insert_before(anchor, f"• {parameter}", "List Paragraph")
    insert_labeled(anchor, "Examples: ", "")
    for example in command["examples"]:
        insert_before(anchor, str(example), "Code")
    insert_labeled(anchor, "Result and consequences: ", str(command["result"]))


def paragraphs_between(document: Document, start_text: str, end_text: str) -> tuple[Paragraph, Paragraph, list[Paragraph]]:
    paragraphs = document.paragraphs
    start_index = next(index for index, p in enumerate(paragraphs) if p.text.strip() == start_text)
    end_index = next(index for index in range(start_index + 1, len(paragraphs)) if paragraphs[index].text.strip() == end_text)
    return paragraphs[start_index], paragraphs[end_index], paragraphs[start_index + 1:end_index]


def replace_w_blocks(document: Document) -> None:
    search_from = 0
    for executable in ("VIC20Emulator.exe", "VIC20EmulatorC.exe"):
        paragraphs = document.paragraphs
        start_index = next(
            index for index in range(search_from, len(paragraphs))
            if paragraphs[index].text.strip() == "/w" and paragraphs[index].style.name == "Heading 1"
        )
        end_index = next(
            index for index in range(start_index + 1, len(paragraphs))
            if paragraphs[index].text.strip() == "/b" and paragraphs[index].style.name == "Heading 1"
        )
        heading = paragraphs[start_index]
        anchor = paragraphs[end_index]
        for paragraph in paragraphs[start_index + 1:end_index]:
            remove_paragraph(paragraph)
        heading.paragraph_format.space_before = Pt(14)
        heading.paragraph_format.keep_with_next = True
        insert_labeled(anchor, "Meaning: ", "Selects the VIC-20 memory-expansion configuration used at startup.")
        insert_labeled(anchor, "General form: ", "")
        insert_before(anchor, "/w[CONF]", "Code")
        configurations = (
            "0 or omitted: unexpanded VIC-20 (default).",
            "1: +3 KB expansion.",
            "2: +8 KB expansion.",
            "3: +16 KB expansion (+8 KB +8 KB).",
            "4: +24 KB expansion (+8 KB +8 KB +8 KB).",
            "5: +32 KB expansion (+8 KB +8 KB +8 KB +8 KB).",
            "6: +11 KB expansion (+3 KB +8 KB).",
            "7: +19 KB expansion (+3 KB +8 KB +8 KB).",
            "8: +27 KB expansion (+3 KB +8 KB +8 KB +8 KB).",
            "9 or any greater numeric value: +35 KB expansion (+3 KB +8 KB +8 KB +8 KB +8 KB); values above 9 are clamped to 9.",
        )
        for configuration in configurations:
            insert_before(anchor, f"• {configuration}", "List Paragraph")
        insert_labeled(anchor, "Examples: ", "")
        for example in (f"{executable} /w0", f"{executable} /w3", f"{executable} /w9"):
            insert_before(anchor, example, "Code")
        insert_labeled(
            anchor,
            "Consequence: ",
            "The selected RAM expansion is installed before the VIC-20 is initialized, changing its visible memory map and the software configurations it can run.",
        )
        search_from = start_index + 1


def remove_spanish_language_entries(document: Document) -> None:
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        if text in {"• ESP: Spanish.", "VIC20Emulator.exe /iESP", "VIC20EmulatorC.exe /iESP"}:
            remove_paragraph(paragraph)


def tailor_inherited_commands(document: Document) -> None:
    paragraphs = document.paragraphs
    start_index = next(
        i for i, p in enumerate(paragraphs)
        if p.text.strip() == "3.2 Inherited Commodore Commands" and p.style.name == "Heading 1"
    )
    end_index = next(
        i for i in range(start_index + 1, len(paragraphs))
        if paragraphs[i].text.strip().startswith("3.3 ") and paragraphs[i].style.name == "Heading 1"
    )
    indices = [i for i in range(start_index + 1, end_index) if paragraphs[i].style.name == "Heading 2"]
    indices.append(end_index)
    for position in range(len(indices) - 1):
        command_index = indices[position]
        command = paragraphs[command_index].text.strip()
        if command not in INHERITED_RESULTS:
            continue
        meaning, result = INHERITED_RESULTS[command]
        for paragraph in paragraphs[command_index + 1:indices[position + 1]]:
            if paragraph.text.startswith("Meaning:"):
                paragraph.clear()
                paragraph.add_run("Meaning: ").bold = True
                paragraph.add_run(meaning)
            elif paragraph.text.startswith("Result and consequences:"):
                paragraph.clear()
                paragraph.add_run("Result and consequences: ").bold = True
                paragraph.add_run(result)


def tailor_standard_grid_commands(document: Document) -> None:
    paragraphs = document.paragraphs
    start_index = next(
        i for i, p in enumerate(paragraphs)
        if p.text.strip() == "GRIDON" and p.style.name == "Heading 2"
    )
    end_index = next(
        i for i in range(start_index + 1, len(paragraphs))
        if paragraphs[i].text.strip() == "GRIDOFF" and paragraphs[i].style.name == "Heading 2"
    )
    block = paragraphs[start_index + 1:end_index]
    examples_index = next(i for i, p in enumerate(block) if p.text.startswith("Examples:"))
    parameters_index = next(i for i, p in enumerate(block) if p.text.startswith("Parameters:"))
    for paragraph in block[parameters_index + 1:examples_index]:
        remove_paragraph(paragraph)
    examples_anchor = next(
        p for p in document.paragraphs[start_index + 1:]
        if p.text.startswith("Examples:")
    )
    insert_before(examples_anchor, "• COLOR: mandatory numeric grid color.", "List Paragraph")

    paragraphs = document.paragraphs
    start_index = next(
        i for i, p in enumerate(paragraphs)
        if p.text.strip() == "GRIDON" and p.style.name == "Heading 2"
    )
    end_index = next(
        i for i in range(start_index + 1, len(paragraphs))
        if paragraphs[i].text.strip() == "GRIDOFF" and paragraphs[i].style.name == "Heading 2"
    )
    for paragraph in paragraphs[start_index + 1:end_index]:
        if paragraph.text.startswith("Meaning:"):
            paragraph.clear()
            paragraph.add_run("Meaning: ").bold = True
            paragraph.add_run("Enables the generic screen grid overlay.")
        elif paragraph.style.name == "Code" and paragraph.text.startswith("GRIDON"):
            paragraph.clear()
            paragraph.add_run("GRIDON COLOR" if " " not in paragraph.text else "GRIDON 14")
        elif paragraph.text.startswith("Result and consequences:"):
            paragraph.clear()
            paragraph.add_run("Result and consequences: ").bold = True
            paragraph.add_run("Enables the screen grid using the requested color.")

    # The first code paragraph is the syntax and the second is the example.
    code_paragraphs = [
        p for p in document.paragraphs[start_index + 1:end_index]
        if p.style.name == "Code"
    ]
    code_paragraphs[0].clear()
    code_paragraphs[0].add_run("GRIDON COLOR")
    code_paragraphs[1].clear()
    code_paragraphs[1].add_run("GRIDON 14")
    for paragraph in code_paragraphs[2:]:
        remove_paragraph(paragraph)

    paragraphs = document.paragraphs
    start_index = next(
        i for i, p in enumerate(paragraphs)
        if p.text.strip() == "GRIDOFF" and p.style.name == "Heading 2"
    )
    end_index = next(
        i for i in range(start_index + 1, len(paragraphs))
        if paragraphs[i].style.name == "Heading 2"
    )
    for paragraph in paragraphs[start_index + 1:end_index]:
        if paragraph.text.startswith("Meaning:"):
            paragraph.clear()
            paragraph.add_run("Meaning: ").bold = True
            paragraph.add_run("Disables the generic screen grid overlay.")
        elif paragraph.text.startswith("Result and consequences:"):
            paragraph.clear()
            paragraph.add_run("Result and consequences: ").bold = True
            paragraph.add_run("Disables the screen grid overlay.")


def replace_specific_section(document: Document) -> None:
    paragraphs = document.paragraphs
    start_index = next(
        i for i, p in enumerate(paragraphs)
        if p.text.strip().startswith("3.3 ") and p.style.name == "Heading 1"
    )
    end_index = next(
        i for i in range(start_index + 1, len(paragraphs))
        if p_text(paragraphs[i]) == "3.4 Local-Console-Only Commands" and paragraphs[i].style.name == "Heading 1"
    )
    anchor = paragraphs[end_index]
    for paragraph in paragraphs[start_index:end_index]:
        remove_paragraph(paragraph)
    section = insert_before(anchor, "3.3 VIC20-Specific Commands", "Heading 1")
    section.paragraph_format.keep_with_next = True
    for command in VIC20_COMMANDS:
        add_command_block(anchor, command)


def p_text(paragraph: Paragraph) -> str:
    return paragraph.text.strip()


def replace_all_text_nodes(document: Document) -> None:
    for node in document.part.element.xpath(".//w:t"):
        if node.text is None:
            continue
        value = node.text
        for old, new in GLOBAL_REPLACEMENTS:
            value = value.replace(old, new)
        node.text = value


def set_update_fields(document: Document) -> None:
    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("destination", type=Path)
    args = parser.parse_args()

    document = Document(args.source)
    replace_all_text_nodes(document)
    remove_spanish_language_entries(document)
    replace_w_blocks(document)
    tailor_standard_grid_commands(document)
    tailor_inherited_commands(document)
    replace_specific_section(document)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        document.styles[style_name].paragraph_format.keep_with_next = True
    document.styles["Heading 2"].paragraph_format.space_before = Pt(14)
    for paragraph in document.paragraphs:
        if paragraph.style.name in {"Title", "Heading 1", "Heading 2", "Heading 3"}:
            paragraph.paragraph_format.keep_with_next = True
        if paragraph.style.name == "Heading 2" or (
            paragraph.style.name == "Heading 1" and paragraph.text.strip().startswith("/")
        ):
            paragraph.paragraph_format.space_before = Pt(14)

    set_update_fields(document)
    document.core_properties.title = "Commodore VIC-20 Emulator User Guide"
    document.core_properties.subject = "Startup options and command console reference"
    args.destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.destination)


if __name__ == "__main__":
    main()
