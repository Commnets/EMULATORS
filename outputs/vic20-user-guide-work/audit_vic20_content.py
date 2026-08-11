from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


import sys

path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(r"C:\Workspaces\EMULATORS\outputs\vic20-user-guide-work\VIC20Emulator_UserGuide.docx")
document = Document(path)
headings = []
for index, paragraph in enumerate(document.paragraphs):
    if paragraph.style and paragraph.style.name in {"Title", "Heading 1", "Heading 2", "Heading 3"}:
        headings.append({"index": index, "style": paragraph.style.name, "text": paragraph.text})

all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
settings = document.settings.element
update = settings.find(qn("w:updateFields"))
result = {
    "paragraphs": len(document.paragraphs),
    "headings": headings,
    "heading2_count": sum(item["style"] == "Heading 2" for item in headings),
    "contains_foreword": "Foreword" in all_text,
    "contains_c64": "C64" in all_text,
    "contains_spanish_language": "ESP: Spanish" in all_text or "/iESP" in all_text,
    "contains_old_specific": any(name in {item["text"] for item in headings} for name in ("CIA1", "CIA2", "PLA", "BITMAPDUMP", "SPRITESDUMP", "SPRITESDRAW", "PADDLE")),
    "update_fields": None if update is None else update.get(qn("w:val")),
    "styles": {
        name: {
            "keep_with_next": document.styles[name].paragraph_format.keep_with_next,
            "space_before": document.styles[name].paragraph_format.space_before.pt if document.styles[name].paragraph_format.space_before else None,
        }
        for name in ("Title", "Heading 1", "Heading 2", "Heading 3")
    },
}
Path(r"C:\Workspaces\EMULATORS\outputs\vic20-user-guide-work\content-audit.json").write_text(
    json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8"
)
