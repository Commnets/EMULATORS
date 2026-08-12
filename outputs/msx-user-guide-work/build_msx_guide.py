from __future__ import annotations

import argparse
import importlib.util
from pathlib import Path

from docx import Document
from docx.shared import Pt


BASE_PATH = Path(r"C:\Workspaces\EMULATORS\outputs\zx81-user-guide-work\build_zx81_guide.py")
SPEC = importlib.util.spec_from_file_location("zx81_guide_support", BASE_PATH)
ZX = importlib.util.module_from_spec(SPEC)
assert SPEC.loader is not None
SPEC.loader.exec_module(ZX)
SUPPORT = ZX.SUPPORT


REPLACEMENTS = (
    ("Sinclair ZX Spectrum Emulator User Guide", "MSX Emulator User Guide"),
    ("ZX Spectrum 16K/48K - Startup and Console Reference", "MSX Models - Startup and Console Reference"),
    ("ZX Spectrum 16K/48K", "MSX Models"),
    ("ZXSpectrumEmulator_UserGuide.docx", "MSXEmulator_UserGuide.docx"),
    ("ZXSpectrumEmulatorC.exe", "MSXEmulatorC.exe"),
    ("ZXSpectrumEmulator.exe", "MSXEmulator.exe"),
    ("ZXSpectrumEmulatorC", "MSXEmulatorC"),
    ("ZXSpectrumEmulator", "MSXEmulator"),
    ("ZXSPECTRUM::ZXSpectrumEmulator", "MSX::MSXEmulator"),
    ("ZXSPECTRUM::CommandBuilder", "MSX::CommandBuilder"),
    ("ZX Spectrum-specific", "MSX-specific"),
    ("ZX Spectrum-Specific", "MSX-Specific"),
    ("ZX Spectrum builder", "MSX builder"),
    ("Sinclair ZX Spectrum", "MSX"),
    ("ZX Spectrum", "MSX"),
    ("zxspectrum.log", "msx.log"),
    ("zxspectrum-screen.png", "msx-screen.png"),
    ("game.tap", "program.cas"),
)


MSX_COMMANDS = (
    {
        "name": "VDP", "aliases": "CVDP.",
        "meaning": "Returns the current state of the MSX video display processor.",
        "syntax": "VDP", "parameters": ["No parameters."], "examples": ["VDP"],
        "result": "Returns the active TMS9918A/TMS9929A VDP information structure, including registers and current video state.",
    },
    {
        "name": "PSG", "aliases": "CPSG.",
        "meaning": "Returns the current state of the MSX programmable sound generator.",
        "syntax": "PSG", "parameters": ["No parameters."], "examples": ["PSG"],
        "result": "Returns the AY-3-8910 PSG information structure and its current register-derived state.",
    },
    {
        "name": "PNAME", "aliases": "CPNAME.",
        "meaning": "Returns a snapshot of the VDP pattern-name table selected by the current VDP registers.",
        "syntax": "PNAME", "parameters": ["No parameters."], "examples": ["PNAME"],
        "result": "Returns the bytes in the active pattern-name table. Their interpretation depends on the current VDP screen mode.",
    },
    {
        "name": "PGEN", "aliases": "CPGEN.",
        "meaning": "Returns a snapshot of the VDP pattern-generator table selected by the current VDP registers.",
        "syntax": "PGEN", "parameters": ["No parameters."], "examples": ["PGEN"],
        "result": "Returns the bytes used by the active VDP pattern-generator table.",
    },
    {
        "name": "CNAME", "aliases": "CCNAME.",
        "meaning": "Returns a snapshot of the VDP colour table selected by the current VDP registers.",
        "syntax": "CNAME", "parameters": ["No parameters."], "examples": ["CNAME"],
        "result": "Returns the bytes in the active colour table. Their interpretation depends on the current VDP screen mode.",
    },
    {
        "name": "SPRITESDRAW", "aliases": "CSPRITESDRAW.",
        "meaning": "Returns textual drawings of all or selected MSX sprite patterns.",
        "syntax": "SPRITESDRAW [1..32 ...]",
        "parameters": ["Sprite numbers: optional values from 1 through 32. Invalid values are ignored and duplicates are coalesced; with none, the command returns all sprites."],
        "examples": ["SPRITESDRAW", "SPRITESDRAW 1 8 16 32"],
        "result": "Returns a textual representation of the selected sprite patterns using the VDP's current sprite tables and mode.",
    },
    {
        "name": "VDPMEMORY", "aliases": "CVDPMEMORY.",
        "meaning": "Reads a byte or inclusive range from VDP video memory.",
        "syntax": "VDPMEMORY ADDRESS [OTHER_ADDRESS]",
        "parameters": [
            "ADDRESS: mandatory first video-memory address.",
            "OTHER_ADDRESS: optional second endpoint. Reversed endpoints are normalized. The parser accepts the Z80's 16-bit address range; meaningful mapped VRAM depends on the active VDP.",
        ],
        "examples": ["VDPMEMORY 0x0000", "VDPMEMORY 0x1800 0x181F"],
        "result": "Returns the addressed VDP-memory bytes; one endpoint reads one byte and two endpoints read the inclusive normalized range.",
    },
    {
        "name": "SETVDPMEMORY", "aliases": "CVDPSETMEMORY.",
        "meaning": "Writes one byte to one VDP-memory address or repeatedly across an inclusive range.",
        "syntax": "SETVDPMEMORY ADDRESS [FINAL_ADDRESS] VALUE",
        "parameters": [
            "ADDRESS: mandatory first video-memory address.",
            "FINAL_ADDRESS: optional range endpoint. Reversed endpoints are normalized.",
            "VALUE: mandatory byte value written at every selected address.",
        ],
        "examples": ["SETVDPMEMORY 0x1800 0x41", "SETVDPMEMORY 0x1800 0x181F 0"],
        "result": "Changes the selected VRAM byte or fills the inclusive range; malformed parameters return an error response.",
    },
    {
        "name": "VDPEVENTS", "aliases": "CVDPEVENTS.",
        "meaning": "Enables or disables visualization of significant CPU/VDP events.",
        "syntax": "VDPEVENTS ON|OFF",
        "parameters": ["ON enables event visualization. Any other single token, including OFF, disables it."],
        "examples": ["VDPEVENTS ON", "VDPEVENTS OFF"],
        "result": "Changes the event-overlay flag used by the active MSX VDP.",
    },
)


STANDARD_OVERRIDES = {
    "HELP": {
        "name": "HELP", "aliases": "? and CHELP.",
        "meaning": "Displays the command help assembled for the MSX emulator.", "syntax": "HELP [COMMAND]",
        "parameters": ["COMMAND: optional command name whose detailed help should be shown; without it, the available help catalogue is displayed."],
        "examples": ["HELP", "HELP VDPMEMORY", "? PSG"],
        "result": "Displays MSX-specific and standard builder help. In the local executable, LocalConsole also supplies help for its eleven emulator/media commands.",
    },
    "SOUNDON": {
        "name": "SOUNDON", "aliases": "CSOUNDON.",
        "meaning": "Enables MSX PSG sound output.", "syntax": "SOUNDON",
        "parameters": ["No parameters."], "examples": ["SOUNDON"],
        "result": "Enables the AY-3-8910 sound device.",
    },
    "SOUNDOFF": {
        "name": "SOUNDOFF", "aliases": "CSOUNDOFF.",
        "meaning": "Disables MSX PSG sound output.", "syntax": "SOUNDOFF",
        "parameters": ["No parameters."], "examples": ["SOUNDOFF"],
        "result": "Disables audible PSG output while emulation continues.",
    },
}


LOCAL_OVERRIDES = {
    "POSSIBLEPERS": {
        "name": "POSSIBLEPERS", "aliases": "None.",
        "meaning": "Lists the peripheral IDs advertised by the MSX I/O peripheral builder.",
        "syntax": "POSSIBLEPERS", "parameters": ["No parameters."], "examples": ["POSSIBLEPERS"],
        "result": "Lists ID 10 (typewriter keyboard injection) and ID 101 (MSX CAS cassette injection).",
    },
    "CONNECTPER": {
        "name": "CONNECTPER", "aliases": "None.",
        "meaning": "Builds and connects an MSX peripheral to the first compatible device.",
        "syntax": "CONNECTPER ID [ATTRS...]",
        "parameters": ["ID: 10 for typewriter input or 101 for CAS cassette injection.", "ATTRS: optional construction attributes; neither currently advertised peripheral requires one."],
        "examples": ["CONNECTPER 10", "CONNECTPER 101"],
        "result": "Connection succeeds only when the ID exists and a compatible target device is available.",
    },
    "LOADPERDATA": {
        "name": "LOADPERDATA", "aliases": "None.",
        "meaning": "Loads keystrokes or cassette media into a connected MSX peripheral.",
        "syntax": "LOADPERDATA ID FILE [DEBUGFILE]",
        "parameters": ["ID: connected peripheral ID.", "FILE: a supported keystroke file for ID 10 or an MSX CAS file for ID 101.", "DEBUGFILE: optional trace file enabled after a successful load."],
        "examples": ["LOADPERDATA 10 startup.txt", "LOADPERDATA 101 program.cas cassette-trace.log"],
        "result": "Loads the file only if the selected peripheral accepts its type. ID 101 handles CAS blocks through the standard MSX cassette BIOS entry points.",
    },
    "EMPTYPERDATA": {
        "name": "EMPTYPERDATA", "aliases": "None.",
        "meaning": "Creates and attaches an empty CAS data container to a connected cassette-injection peripheral.",
        "syntax": "EMPTYPERDATA ID FILE",
        "parameters": ["ID: normally 101.", "FILE: destination path; use the .cas extension."],
        "examples": ["EMPTYPERDATA 101 recording.cas"],
        "result": "Prepares empty writable cassette data. The emulated BIOS TAPOON/TAPOUT/TAPOOF calls can then record blocks into it.",
    },
    "SAVEPERDATA": {
        "name": "SAVEPERDATA", "aliases": "None.",
        "meaning": "Saves the current data retrieved from a connected peripheral.",
        "syntax": "SAVEPERDATA ID [FILE]",
        "parameters": ["ID: connected peripheral ID, normally 101 for cassette data.", "FILE: optional output path; it may be omitted after LOADPERDATA or EMPTYPERDATA established a filename."],
        "examples": ["SAVEPERDATA 101", "SAVEPERDATA 101 copy.cas"],
        "result": "Writes the cassette's current CAS blocks when the format is writable; otherwise an error is returned.",
    },
}


def replace_text_nodes(document: Document) -> None:
    for node in document.part.element.xpath(".//w:t"):
        if node.text is None:
            continue
        value = node.text
        for old, new in REPLACEMENTS:
            value = value.replace(old, new)
        node.text = value


def executable(occurrence: int) -> str:
    return "MSXEmulator.exe" if occurrence == 0 else "MSXEmulatorC.exe"


def add_remote(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Enables or disables the TCP remote-command system in the executable without a local console.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/o[YES|NO]", "Code")
    ZX.bullet(anchor, "YES or an empty value: create and start the remote command system.")
    ZX.bullet(anchor, "NO (and any other non-YES non-empty value): leave it disabled.")
    ZX.bullet(anchor, "Omitted: no remote listener is created.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, "MSXEmulator.exe /o", "Code")
    SUPPORT.insert_before(anchor, "MSXEmulator.exe /oYES /p61000", "Code")
    SUPPORT.insert_before(anchor, "MSXEmulator.exe /oNO", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "When enabled, remote clients can execute the MSX and standard builder commands. LocalConsole-only commands remain unavailable.")


def add_language(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Selects the ROM/keyboard language used by the selected MSX model.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/iLANG", "Code")
    for text in (
        "SVI728: ENG (default) or ESP; other values make ROM loading fail.",
        "SONYHB10P: always uses its English ROM; the supplied value is ignored.",
        "PHILIPSVG8010: FRA selects French; ENG and unrecognized values use English.",
        "CANONV20: always uses its English ROM; the supplied value is ignored.",
        "SVI738: the code accepts ENG, DEU and SEW. The help banner says SWE, but that spelling is not accepted; SVI738 is currently incomplete and must not be selected for normal use.",
        "MSXSTD/default: always uses the English SVI728-compatible ROM.",
        "Language identifiers are case-sensitive; use the uppercase forms above.",
    ):
        ZX.bullet(anchor, text)
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /mSVI728 /iESP", "Code")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /mPHILIPSVG8010 /iFRA", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "The selected model loads its matching BASIC/BIOS ROM before initialization. The required file must exist under ./bios relative to the process working directory.")


def add_sound_startup(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Selects the initial AY-3-8910 PSG output state.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/r[YES|NO]", "Code")
    ZX.bullet(anchor, "YES or an empty value: sound enabled.")
    ZX.bullet(anchor, "NO (and any other non-YES non-empty value): sound disabled.")
    ZX.bullet(anchor, "Omitted: sound enabled.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /r", "Code")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /rNO", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "The active model's PSG starts enabled or disabled as requested.")


def add_ntsc(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Requests NTSC timing and the TMS9918A instead of PAL timing and the TMS9929A.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/n", "Code")
    ZX.bullet(anchor, "No value: presence selects NTSC for each recognized named model; omission selects PAL.")
    ZX.bullet(anchor, "The default MSXSTD model is fixed to PAL, so /n has no effect when /m is omitted or unknown.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /mSVI728 /n", "Code")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /mCANONV20", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "A recognized named model uses the requested PAL/NTSC VDP and geometry. The MSX CPU clock remains 3.58 MHz in the implemented MSX1 models.")


def add_border(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Draws the diagnostic grid over the writable screen area and selects its startup colour.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/b[COLOR]", "Code")
    ZX.bullet(anchor, "COLOR omitted or 0: colour 0.")
    ZX.bullet(anchor, "1 through 3: requested startup grid colour.")
    ZX.bullet(anchor, "Any value greater than 3: clamped to 3 by the startup parser.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /b", "Code")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /b3", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "The MSX screen starts with its diagnostic grid enabled in the selected colour.")


def add_configuration(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Requests a model memory configuration.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/wCONF", "Code")
    ZX.bullet(anchor, "0 or omitted: the only configuration currently admitted by every defined model.")
    ZX.bullet(anchor, "Other values: rejected by the selected model and adjusted back to configuration 0.")
    ZX.bullet(anchor, "Configuration 0 provides 16 KiB RAM for MSXSTD, SVI728, SONYHB10P and the incomplete SVI738; 32 KiB for PHILIPSVG8010; and 64 KiB for CANONV20.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /mSVI728 /w0", "Code")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /mCANONV20 /w0", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "The model validates the request and uses its only basic mapping. /w does not currently change RAM capacity; /m selects the model and therefore the installed RAM.")


def add_machine(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Selects the MSX machine model and its ROM, memory map and hardware profile.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/mMACHINE", "Code")
    for text in (
        "SVI728: Spectravideo SVI-728 MSX1, 16 KiB RAM; English or Spanish ROM.",
        "SONYHB10P: Sony HB-10P MSX1, 16 KiB RAM; English ROM.",
        "PHILIPSVG8010: Philips VG-8010 MSX1, 32 KiB RAM; English or French ROM.",
        "CANONV20: Canon V-20 MSX1, 64 KiB RAM; English ROM.",
        "SVI738: advertised as Spectravideo SVI-738/MSX2 with 16 KiB RAM, but its VDP and PSG factories return null and the implementation is not operational.",
        "Omitted or unknown: MSXSTD, a PAL-only standard MSX1 model equivalent to an English SVI728 with 16 KiB RAM. Unknown names are logged.",
    ):
        ZX.bullet(anchor, text)
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /mSVI728 /iESP", "Code")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /mCANONV20 /n", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "The selected model constructs its corresponding memory, VDP, PSG and ROM profile. Do not select SVI738 until its missing devices are implemented.")


def add_port_or_peripherals(anchor, occurrence: int) -> None:
    if occurrence == 0:
        SUPPORT.insert_labeled(anchor, "Meaning: ", "Selects the TCP listening port used by the non-console executable's remote command channel.")
        SUPPORT.insert_labeled(anchor, "General form: ", "")
        SUPPORT.insert_before(anchor, "/pPORTNUMBER", "Code")
        ZX.bullet(anchor, "0 through 65535: requested listening port; values above 1000 are recommended.")
        ZX.bullet(anchor, "Omitted: port 60000 is used when /o enables the channel.")
        ZX.bullet(anchor, "Collision: the inherited emulator also interprets /p as a comma-separated peripheral-ID list. A supplied port number is therefore also attempted as a peripheral ID during initialization.")
        SUPPORT.insert_labeled(anchor, "Examples: ", "")
        SUPPORT.insert_before(anchor, "MSXEmulator.exe /o", "Code")
        SUPPORT.insert_before(anchor, "MSXEmulator.exe /o /p61000", "Code")
        SUPPORT.insert_labeled(anchor, "Consequence: ", "The remote server listens on the selected port. When /p is present, the same value is also attempted as a peripheral ID and normally produces a log warning.")
    else:
        SUPPORT.insert_labeled(anchor, "Meaning: ", "Connects one or more peripherals during startup.")
        SUPPORT.insert_labeled(anchor, "General form: ", "")
        SUPPORT.insert_before(anchor, "/pPER1[,PER2...]", "Code")
        ZX.bullet(anchor, "10: typewriter keyboard injection.")
        ZX.bullet(anchor, "101: MSX CAS cassette injection through standard BIOS traps.")
        SUPPORT.insert_labeled(anchor, "Examples: ", "")
        SUPPORT.insert_before(anchor, "MSXEmulatorC.exe /p10", "Code")
        SUPPORT.insert_before(anchor, "MSXEmulatorC.exe /p10,101", "Code")
        SUPPORT.insert_labeled(anchor, "Consequence: ", "Each requested peripheral is built and connected to the first compatible device. Unsupported IDs are logged and initialization continues.")


def replace_family_sections(document: Document) -> None:
    paragraphs = document.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith("3.2 ") and p.style.name == "Heading 1")
    end = next(i for i in range(start + 1, len(paragraphs)) if paragraphs[i].text.strip() == "3.4 Local-Console-Only Commands" and paragraphs[i].style.name == "Heading 1")
    anchor = paragraphs[end]
    for paragraph in paragraphs[start:end]:
        SUPPORT.remove_paragraph(paragraph)
    heading = SUPPORT.insert_before(anchor, "3.2 MSX-Specific Commands", "Heading 1")
    heading.paragraph_format.keep_with_next = True
    for command in MSX_COMMANDS:
        ZX.add_command_block(anchor, command)
    anchor.text = "3.3 Local-Console-Only Commands"


def replace_narrative(document: Document) -> None:
    replacements = {
        "This guide describes the startup and command interfaces": "This guide describes the startup and command interfaces implemented by the two MSX executables in EMULATORS: MSXEmulator (without a local command console) and MSXEmulatorC (with a local command console). It covers the implemented MSX model selectors, PAL and NTSC operation, model-dependent ROM languages, and is based on both entry points, the MSX::MSXEmulator inheritance chain, LocalConsole, and the effective MSX::CommandBuilder -> MCHEmul::StandardCommandBuilder -> MCHEmul::StandardCommandBuilder chain constructed by the entry points.",
        "The non-console executable does not expose": "The non-console executable does not expose an interactive prompt. When started with /o it creates a remote communication system using the MSX and standard command builders. Commands marked Local console only are intercepted by LocalConsole and therefore cannot be sent to the non-console executable.",
        "Executable: MSXEmulator.exe.": "Executable: MSXEmulator.exe. It accepts the framework and MSX-specific options below, plus /o and the communication-port option. Option letters and language identifiers are case-sensitive because the command-line parser preserves them exactly.",
        "Executable: MSXEmulatorC.exe.": "Executable: MSXEmulatorC.exe. Its main program delegates the complete command line to MSX::MSXEmulator, so all inherited framework and MSX-specific options are active even though the short main-level banner only advertises /h. It has no /o listener; commands are entered in its local console.",
        "The local console checks its eleven": "The local console checks its eleven emulator-level commands first. All remaining text is delegated through MSX::CommandBuilder and the standard builder chain. The current entry points pass a StandardCommandBuilder into MSX::CommandBuilder, whose constructor adds another StandardCommandBuilder parent; duplicate registrations resolve to the same standard command set. The remote /o channel in MSXEmulator.exe uses the same builder chain.",
        "The table distinguishes the interactive console": "The table distinguishes the interactive console from the command channel of the non-console executable. Remote availability assumes MSXEmulator.exe was started with /o and that the client is connected.",
        "This file is a maintained repository artifact.": "This file is a maintained repository artifact. Every change that adds, removes, renames or changes an MSX startup option, model selector, command-builder registration, LocalConsole-only command, syntax, parameter meaning, consequence or visible formatted output must include a review and, when applicable, an update of docs/MSXData/MSXEmulator_UserGuide.docx. Update the corresponding canonical .fmt source whenever the command's InfoStructure contract changes.",
        "Audit the complete chain, not only the edited class:": "Audit the complete chain, not only the edited class: both MSXEmulator entry points and the MSX::MSXEmulator hierarchy for startup options; LocalConsole for local-only commands; and the complete effective MSX::CommandBuilder -> MCHEmul::StandardCommandBuilder -> MCHEmul::StandardCommandBuilder chain for builder commands.",
    }
    for paragraph in document.paragraphs:
        for prefix, value in replacements.items():
            if paragraph.text.startswith(prefix):
                paragraph.clear()
                paragraph.add_run(value)
                break


def update_summary_table(document: Document) -> None:
    table = document.tables[0]
    table.cell(0, 1).text = "MSXEmulatorC local console"
    table.cell(0, 2).text = "MSXEmulator remote /o"
    table.cell(1, 0).text = "Standard builder"
    table.cell(2, 0).text = "MSX builder"
    table.cell(3, 0).text = "LocalConsole emulator/media operations"
    table.cell(3, 1).text = "Yes"
    table.cell(3, 2).text = "No"
    table._tbl.remove(table.rows[4]._tr)


def correct_console_examples(document: Document) -> None:
    paragraphs = document.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "2. Startup Parameters: Version With a Local Console" and p.style.name == "Title")
    end = next(i for i in range(start + 1, len(paragraphs)) if paragraphs[i].text.strip() == "3. Command Console Reference" and paragraphs[i].style.name == "Title")
    for paragraph in paragraphs[start:end]:
        for run in paragraph.runs:
            run.text = run.text.replace("MSXEmulator.exe", "MSXEmulatorC.exe")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("destination", type=Path)
    args = parser.parse_args()

    document = Document(args.source)
    replace_text_nodes(document)
    ZX.replace_heading_body(document, "/o", "Heading 1", add_remote)
    ZX.replace_heading_body(document, "/i", "Heading 1", add_language)
    ZX.replace_heading_body(document, "/r", "Heading 1", add_sound_startup)
    ZX.replace_heading_body(document, "/p", "Heading 1", add_port_or_peripherals)
    ZX.replace_heading_body(document, "/n", "Heading 1", add_ntsc)
    ZX.replace_heading_body(document, "/b", "Heading 1", add_border)
    ZX.replace_heading_body(document, "/w", "Heading 1", add_configuration)
    ZX.replace_heading_body(document, "/m", "Heading 1", add_machine)
    replace_family_sections(document)
    for name, command in STANDARD_OVERRIDES.items():
        ZX.replace_standard_command(document, name, command)
    for name, command in LOCAL_OVERRIDES.items():
        ZX.replace_standard_command(document, name, command, "Local console only.")
    replace_narrative(document)
    update_summary_table(document)
    correct_console_examples(document)
    ZX.normalize_bullets(document)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        document.styles[style_name].paragraph_format.keep_with_next = True
    document.styles["Heading 2"].paragraph_format.space_before = Pt(14)
    for paragraph in document.paragraphs:
        if paragraph.style.name in {"Title", "Heading 1", "Heading 2", "Heading 3"}:
            paragraph.paragraph_format.keep_with_next = True
        if paragraph.style.name == "Heading 2" or (paragraph.style.name == "Heading 1" and paragraph.text.strip().startswith("/")):
            paragraph.paragraph_format.space_before = Pt(14)

    ZX.set_update_fields(document)
    document.core_properties.title = "MSX Emulator User Guide"
    document.core_properties.subject = "MSX models: startup options and command console reference"
    args.destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.destination)


if __name__ == "__main__":
    main()
