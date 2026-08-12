from __future__ import annotations

import argparse
import importlib.util
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SUPPORT_PATH = Path(r"C:\Workspaces\EMULATORS\outputs\c264-user-guide-work\build_c264_guide.py")
SPEC = importlib.util.spec_from_file_location("c264_guide_support", SUPPORT_PATH)
SUPPORT = importlib.util.module_from_spec(SPEC)
assert SPEC.loader is not None
SPEC.loader.exec_module(SUPPORT)


REPLACEMENTS = (
    ("Commodore 264 Series Emulator User Guide", "Sinclair ZX80/ZX81 Emulator User Guide"),
    ("C16, C116 and Plus/4 Startup and Console Reference", "ZX80 and ZX81 ROM Variants - Startup and Console Reference"),
    ("C16, C116 and Plus/4", "ZX80 and ZX81 ROM Variants"),
    ("C264Emulator_UserGuide.docx", "ZX81Emulator_UserGuide.docx"),
    ("C264EmulatorC.exe", "ZX81EmulatorC.exe"),
    ("C264Emulator.exe", "ZX81Emulator.exe"),
    ("C264EmulatorC", "ZX81EmulatorC"),
    ("C264Emulator", "ZX81Emulator"),
    ("C264::C264Emulator", "ZX81::ZX81Emulator"),
    ("C264::CommandBuilder", "ZX81::CommandBuilder"),
    ("COMMODORE::CommandBuilder", "SINCLAIR::CommandBuilder"),
    ("Commodore 264 Series", "Sinclair ZX80/ZX81"),
    ("C264-series", "ZX80/ZX81"),
    ("C264 series", "ZX80/ZX81"),
    ("C264-specific", "ZX80/ZX81-specific"),
    ("C264-Specific", "ZX80/ZX81-Specific"),
    ("C264 builder", "ZX80/ZX81 builder"),
    ("Commodore", "Sinclair"),
    ("7501", "Z80"),
    ("c264.log", "zx81.log"),
    ("c264-screen.png", "zx81-screen.png"),
    ("games/demo.prg", "games/demo.bin"),
    ("program.prg", "program.bin"),
    ("demo.prg", "demo.bin"),
)


SINCLAIR_COMMANDS = (
    {
        "name": "SYSVARS",
        "aliases": "CSYSVARS.",
        "meaning": "Returns every system variable defined for the selected ZX80 or ZX81 ROM profile.",
        "syntax": "SYSVARS",
        "parameters": ["No parameters."],
        "examples": ["SYSVARS"],
        "result": "Returns the current values through the system-variable definitions loaded for the selected ROM variant.",
    },
    {
        "name": "SYSVAR",
        "aliases": "CSYSVAR.",
        "meaning": "Returns one named system variable.",
        "syntax": "SYSVAR VNAME",
        "parameters": ["VNAME: mandatory variable name from the selected machine's system-variable table."],
        "examples": ["SYSVAR D_FILE", "SYSVAR E_LINE"],
        "result": "Returns the named variable and its current memory value; an unknown name returns an error entry.",
    },
)


ZX81_COMMANDS = (
    {
        "name": "ULA",
        "aliases": "CULA.",
        "meaning": "Returns the current ULA and video-generation state.",
        "syntax": "ULA",
        "parameters": ["No parameters."],
        "examples": ["ULA"],
        "result": "Returns the ULA state for the selected ZX80 or ZX81 ROM variant, including its current raster and display-generation data.",
    },
    {
        "name": "ULAEVENTS",
        "aliases": "CULAEVENTS.",
        "meaning": "Enables or disables visualization of significant CPU/ULA events.",
        "syntax": "ULAEVENTS ON|OFF",
        "parameters": ["ON: enables event visualization.", "OFF: disables it. The implementation accepts any single token; only the exact value ON enables the feature."],
        "examples": ["ULAEVENTS ON", "ULAEVENTS OFF"],
        "result": "Changes the ULA event-overlay flag used by the active ZX80/ZX81 screen implementation.",
    },
    {
        "name": "DFDUMP",
        "aliases": "CDFDUMP.",
        "meaning": "Returns a snapshot of the current display file.",
        "syntax": "DFDUMP",
        "parameters": ["No parameters."],
        "examples": ["DFDUMP"],
        "result": "Returns the bytes selected by the active ZX80/ZX81 display-file mapping.",
    },
    {
        "name": "CHARSDRAW",
        "aliases": "CCHARSDRAW.",
        "meaning": "Returns textual drawings of all or selected character glyphs.",
        "syntax": "CHARSDRAW [0..63 ...]",
        "parameters": ["Character codes: optional values from 0 through 63; out-of-range values are ignored and duplicates are coalesced. With none, all 64 glyphs are rendered."],
        "examples": ["CHARSDRAW", "CHARSDRAW 0 1 32 63"],
        "result": "Returns textual glyph drawings from the character data visible to the selected ZX80/ZX81 model.",
    },
)


STANDARD_OVERRIDES = {
    "SOUNDON": {
        "name": "SOUNDON", "aliases": "CSOUNDON.",
        "meaning": "Requests that sound output be enabled.", "syntax": "SOUNDON",
        "parameters": ["No parameters."], "examples": ["SOUNDON"],
        "result": "The command is accepted, but the implemented ZX80/ZX81 computer has no sound device, so it produces no state change.",
    },
    "SOUNDOFF": {
        "name": "SOUNDOFF", "aliases": "CSOUNDOFF.",
        "meaning": "Requests that sound output be disabled.", "syntax": "SOUNDOFF",
        "parameters": ["No parameters."], "examples": ["SOUNDOFF"],
        "result": "The command is accepted, but the implemented ZX80/ZX81 computer has no sound device, so it produces no state change.",
    },
    "GRIDON": {
        "name": "GRIDON", "aliases": "CGRIDON.",
        "meaning": "Enables the diagnostic grid overlay on the screen.", "syntax": "GRIDON COLOR",
        "parameters": ["COLOR: mandatory numeric grid color."], "examples": ["GRIDON 14"],
        "result": "Enables the standard screen grid with the requested color.",
    },
    "GRIDOFF": {
        "name": "GRIDOFF", "aliases": "CGRIDOFF.",
        "meaning": "Disables the diagnostic grid overlay.", "syntax": "GRIDOFF",
        "parameters": ["No parameters."], "examples": ["GRIDOFF"],
        "result": "Disables the standard screen grid.",
    },
}


LOCAL_OVERRIDES = {
    "POSSIBLEPERS": {
        "name": "POSSIBLEPERS", "aliases": "None.",
        "meaning": "Lists the peripheral IDs and construction attributes advertised by the active ZX80/ZX81 I/O peripheral builder.",
        "syntax": "POSSIBLEPERS", "parameters": ["No parameters."],
        "examples": ["POSSIBLEPERS"],
        "result": "Lists IDs 10 (typewriter), 100 (cassette), 101 (cassette injection) and 102 (thermal printer). Cartridge ID 201 is accepted by the builder but is not included in this advertised list.",
    },
    "CONNECTPER": {
        "name": "CONNECTPER", "aliases": "None.",
        "meaning": "Builds and connects a peripheral to the first compatible ZX80/ZX81 device.",
        "syntax": "CONNECTPER ID [ATTRS...]",
        "parameters": [
            "ID: peripheral-emulation ID (10, 100, 101, 102 or the additionally supported cartridge ID 201).",
            "ATTRS: optional construction attributes. For printer 102, use f:FILENAME and p:THERMAL or p:THERMAL-PS.",
        ],
        "examples": ["CONNECTPER 100", "CONNECTPER 102 f:printer.ps p:THERMAL-PS"],
        "result": "Connection succeeds only when the ID, attributes and a compatible target device are available; otherwise the answer reports an error.",
    },
    "LOADPERDATA": {
        "name": "LOADPERDATA", "aliases": "None.",
        "meaning": "Loads media or program data into a connected peripheral.",
        "syntax": "LOADPERDATA ID FILE [DEBUGFILE]",
        "parameters": [
            "ID: connected peripheral ID.",
            "FILE: supported data file; for example sampled TZX data for ID 100 or an O/P program for injection ID 101.",
            "DEBUGFILE: optional file that starts full-range format-related debugging after a successful load.",
        ],
        "examples": ["LOADPERDATA 100 tape.tzx", "LOADPERDATA 101 game.p load-trace.log"],
        "result": "The operation fails if the peripheral does not accept the supplied file type; otherwise the data is attached and optional debugging begins.",
    },
    "EMPTYPERDATA": {
        "name": "EMPTYPERDATA", "aliases": "None.",
        "meaning": "Creates and attaches empty media data when the peripheral implements that operation.",
        "syntax": "EMPTYPERDATA ID FILE",
        "parameters": ["ID: connected peripheral ID.", "FILE: destination whose extension selects a writable format."],
        "examples": ["EMPTYPERDATA 100 blank.tzx"],
        "result": "The Sinclair cassette implementation can create an empty TZX container; unsupported peripherals or formats return an error.",
    },
    "SAVEPERDATA": {
        "name": "SAVEPERDATA", "aliases": "None.",
        "meaning": "Saves a peripheral's current media data.",
        "syntax": "SAVEPERDATA ID [FILE]",
        "parameters": [
            "ID: connected peripheral ID.",
            "FILE: optional output path; it may be omitted when LOADPERDATA or EMPTYPERDATA established the current filename.",
        ],
        "examples": ["SAVEPERDATA 100", "SAVEPERDATA 100 copy.tzx"],
        "result": "Saves the current cassette data when the peripheral can retrieve it and the selected file format can be written.",
    },
}


def replace_text_nodes(document: Document) -> None:
    for node in document.part.element.xpath(".//w:t"):
        if node.text is None:
            continue
        value = node.text
        for old, new in REPLACEMENTS:
            value = value.replace(old, new)
        node.text = value


def apply_bullet_numbering(paragraph) -> None:
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "12")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    paragraph._p.get_or_add_pPr().append(num_pr)


def bullet(anchor, text: str) -> None:
    paragraph = SUPPORT.insert_before(anchor, text, "List Paragraph")
    apply_bullet_numbering(paragraph)


def normalize_bullets(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text
        if len(text) >= 2 and text[0] in {"•", "�"} and text[1] == " ":
            paragraph.clear()
            paragraph.add_run(text[2:])
            paragraph.style = "List Paragraph"
            apply_bullet_numbering(paragraph)


def add_reference_body(anchor, command: dict[str, object], availability: str) -> None:
    SUPPORT.insert_labeled(anchor, "Availability: ", availability)
    SUPPORT.insert_labeled(anchor, "Accepted aliases: ", str(command["aliases"]))
    SUPPORT.insert_labeled(anchor, "Meaning: ", str(command["meaning"]))
    SUPPORT.insert_labeled(anchor, "General syntax: ", "")
    SUPPORT.insert_before(anchor, str(command["syntax"]), "Code")
    SUPPORT.insert_labeled(anchor, "Parameters: ", "")
    for parameter in command["parameters"]:
        bullet(anchor, str(parameter))
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    for example in command["examples"]:
        SUPPORT.insert_before(anchor, str(example), "Code")
    SUPPORT.insert_labeled(anchor, "Result and consequences: ", str(command["result"]))


def add_command_block(anchor, command: dict[str, object], availability: str = "Local console and remote /o channel.") -> None:
    heading = SUPPORT.insert_before(anchor, str(command["name"]), "Heading 2")
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.keep_with_next = True
    add_reference_body(anchor, command, availability)


def replace_heading_body(document: Document, name: str, style: str, builder) -> None:
    occurrence = 0
    while True:
        paragraphs = document.paragraphs
        starts = [i for i, p in enumerate(paragraphs) if p.text.strip() == name and p.style.name == style]
        if occurrence >= len(starts):
            break
        start = starts[occurrence]
        heading = paragraphs[start]
        end = next(i for i in range(start + 1, len(paragraphs)) if paragraphs[i].style.name in {style, "Heading 1", "Title"})
        anchor = paragraphs[end]
        for paragraph in paragraphs[start + 1:end]:
            SUPPORT.remove_paragraph(paragraph)
        heading.paragraph_format.space_before = Pt(14)
        heading.paragraph_format.keep_with_next = True
        builder(anchor, occurrence)
        occurrence += 1


def executable(occurrence: int) -> str:
    return "ZX81Emulator.exe" if occurrence == 0 else "ZX81EmulatorC.exe"


def add_language(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Selects the language identifier exposed to the emulated machine.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/iLANG", "Code")
    bullet(anchor, "ENG: English; this is the only language advertised and supported by the ZX80/ZX81 emulator.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /iENG", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "The machine starts with the English language identifier. Other values are passed through but have no documented ZX80/ZX81 ROM variant in this implementation.")


def add_ntsc(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Requests NTSC visualization parameters.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/n", "Code")
    bullet(anchor, "No value: the presence of the option requests NTSC.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /mZX813 /n", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "The current factory always constructs PAL ZX80 and ZX81 machines. The option is detected and a limitation is written to the log, but PAL timing and screen geometry remain active.")


def add_border(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Draws a diagnostic grid around the writable screen area and selects its color.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/b[COLOR]", "Code")
    bullet(anchor, "COLOR omitted or 0: black (default).")
    bullet(anchor, "1 through 15: the requested grid color.")
    bullet(anchor, "Any value greater than 15: clamped to 15.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /b", "Code")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /b14", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "The screen is initialized with the diagnostic grid enabled in the selected color.")


def add_configuration(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Selects the startup RAM configuration.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/wCONF", "Code")
    bullet(anchor, "ZX80: only configuration 0 (unexpanded) is implemented; another value is ignored and logged.")
    bullet(anchor, "ZX81 ROM variants, 0 or omitted: unexpanded configuration (default).")
    bullet(anchor, "ZX81 ROM variants, 1: 16 KiB expansion.")
    bullet(anchor, "ZX81 ROM variants, any value other than 0 or 1: reset to unexpanded configuration and logged.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /mZX80 /w0", "Code")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /mZX813 /w1", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "The selected memory map is installed before the ROM and machine initialize. The broader 3K/8K/24K list in an old header comment is not implemented by the current factory.")


def add_sound_startup(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Requests the initial sound-output state.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/r[YES|NO]", "Code")
    bullet(anchor, "YES or an empty value: requests sound enabled.")
    bullet(anchor, "NO (and any other non-YES non-empty value): requests sound disabled.")
    bullet(anchor, "Omitted: requests sound enabled.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /r", "Code")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /rNO", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "The option is parsed, but the implemented ZX80/ZX81 computer has no sound device, so initialization makes no audio-state change.")


def add_machine(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Selects the ZX80 or ZX81 ROM variant to emulate.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/mMACHINE", "Code")
    bullet(anchor, "ZX80: ZX80 model and ROM; this is the default when /m is omitted.")
    bullet(anchor, "ZX811: ZX81 with the old ROM image.")
    bullet(anchor, "ZX812: ZX81 with the rare/intermediate ROM image.")
    bullet(anchor, "ZX813: ZX81 with the newest ROM image.")
    bullet(anchor, "Any other value: logs the unsupported selector and falls back to ZX80.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /mZX80", "Code")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /mZX811", "Code")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /mZX813 /w1", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "The selector chooses the ROM image and machine type created before initialization. It does not independently select a ULA silicon revision, video standard or RAM expansion.")


def add_port_or_peripherals(anchor, occurrence: int) -> None:
    if occurrence == 0:
        SUPPORT.insert_labeled(anchor, "Meaning: ", "Selects the TCP listening port used by the non-console executable's remote command channel.")
        SUPPORT.insert_labeled(anchor, "General form: ", "")
        SUPPORT.insert_before(anchor, "/pPORTNUMBER", "Code")
        bullet(anchor, "0 through 65535: requested listening port; values above 1000 are recommended.")
        bullet(anchor, "Omitted: port 60000 is used when /o enables the channel.")
        bullet(anchor, "Collision: the inherited emulator also interprets /p as a comma-separated peripheral-ID list. A supplied port number is therefore also attempted as a peripheral ID during initialization.")
        SUPPORT.insert_labeled(anchor, "Examples: ", "")
        SUPPORT.insert_before(anchor, "ZX81Emulator.exe /o", "Code")
        SUPPORT.insert_before(anchor, "ZX81Emulator.exe /o /p61000", "Code")
        SUPPORT.insert_labeled(anchor, "Consequence: ", "The remote server listens on the selected port. When /p is present, the same value is also passed to peripheral startup and normally produces a log warning because a port number is not a valid peripheral ID.")
    else:
        SUPPORT.insert_labeled(anchor, "Meaning: ", "Connects one or more peripherals during startup.")
        SUPPORT.insert_labeled(anchor, "General form: ", "")
        SUPPORT.insert_before(anchor, "/pPER1[,PER2...]", "Code")
        for value in (
            "10: typewriter keyboard input.",
            "100: Sinclair cassette signal simulation.",
            "101: ZX80/ZX81 cassette-data injection.",
            "102: ZX Printer thermal-printer simulation.",
            "201: cartridge implementation supported by the builder, although it is not listed by POSSIBLEPERS.",
        ):
            bullet(anchor, value)
        SUPPORT.insert_labeled(anchor, "Examples: ", "")
        SUPPORT.insert_before(anchor, "ZX81EmulatorC.exe /p10", "Code")
        SUPPORT.insert_before(anchor, "ZX81EmulatorC.exe /p100,102", "Code")
        SUPPORT.insert_labeled(anchor, "Consequence: ", "Each requested peripheral is built and connected to the first compatible device during emulator initialization. Unsupported IDs are logged and initialization continues.")


def replace_family_sections(document: Document) -> None:
    paragraphs = document.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith("3.2 ") and p.style.name == "Heading 1")
    end = next(i for i in range(start + 1, len(paragraphs)) if paragraphs[i].text.strip() == "3.4 Local-Console-Only Commands" and paragraphs[i].style.name == "Heading 1")
    anchor = paragraphs[end]
    for paragraph in paragraphs[start:end]:
        SUPPORT.remove_paragraph(paragraph)
    heading = SUPPORT.insert_before(anchor, "3.2 Inherited Sinclair Commands", "Heading 1")
    heading.paragraph_format.keep_with_next = True
    for command in SINCLAIR_COMMANDS:
        add_command_block(anchor, command)
    heading = SUPPORT.insert_before(anchor, "3.3 ZX80/ZX81-Specific Commands", "Heading 1")
    heading.paragraph_format.keep_with_next = True
    for command in ZX81_COMMANDS:
        add_command_block(anchor, command)


def replace_standard_command(document: Document, name: str, command: dict[str, object], availability: str = "Local console and remote /o channel.") -> None:
    paragraphs = document.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == name and p.style.name == "Heading 2")
    end = next(i for i in range(start + 1, len(paragraphs)) if paragraphs[i].style.name in {"Heading 2", "Heading 1", "Title"})
    anchor = paragraphs[end]
    for paragraph in paragraphs[start + 1:end]:
        SUPPORT.remove_paragraph(paragraph)
    add_reference_body(anchor, command, availability)


def replace_narrative(document: Document) -> None:
    replacements = {
        "This guide describes the startup and command interfaces": "This guide describes the startup and command interfaces implemented by the two Sinclair ZX80/ZX81 executables in EMULATORS: ZX81Emulator (without a local command console) and ZX81EmulatorC (with a local command console). It covers the ZX80 and the three ZX81 ROM selectors and is based on both executable entry points, the ZX81::ZX81Emulator inheritance chain, LocalConsole, and the complete ZX81 -> SINCLAIR -> StandardCommandBuilder responsibility chain.",
        "Executable: ZX81Emulator.exe.": "Executable: ZX81Emulator.exe. It accepts the framework and ZX80/ZX81-specific options below, plus /o and the communication-port option. Option letters are case-sensitive because the command-line parser stores the character exactly as supplied.",
        "Executable: ZX81EmulatorC.exe.": "Executable: ZX81EmulatorC.exe. Its main program delegates the complete command line to ZX81::ZX81Emulator, so all inherited framework options and ZX80/ZX81-specific options are active even though the short main-level help banner only advertises /h. It has no /o listener option; commands are entered in its local console.",
        "The local console checks its eleven": "The local console checks its eleven emulator-level commands first. All remaining text is delegated through ZX81::CommandBuilder, then SINCLAIR::CommandBuilder, and finally MCHEmul::StandardCommandBuilder. The remote /o channel in ZX81Emulator.exe uses only that three-builder chain. The availability line in each entry makes this boundary explicit.",
        "This file is a maintained repository artifact.": "This file is a maintained repository artifact. Every change that adds, removes, renames or changes a ZX80/ZX81 startup option, a command-builder registration, a LocalConsole-only command, command syntax, parameter meaning, command consequences or user-visible formatted output must include a review and, when applicable, an update of docs/ZX81Data/ZX81Emulator_UserGuide.docx. The corresponding canonical .fmt source must be updated whenever the command's InfoStructure contract changes.",
        "Audit the complete chain, not only the edited class:": "Audit the complete chain, not only the edited class: the ZX81Emulator entry point and ZX81::ZX81Emulator hierarchy for startup options; LocalConsole for local-only commands; and ZX81::CommandBuilder -> SINCLAIR::CommandBuilder -> MCHEmul::StandardCommandBuilder for builder commands.",
    }
    for paragraph in document.paragraphs:
        for prefix, value in replacements.items():
            if paragraph.text.startswith(prefix):
                paragraph.clear()
                paragraph.add_run(value)
                break


def correct_console_examples(document: Document) -> None:
    paragraphs = document.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "2. Startup Parameters: Version With a Local Console" and p.style.name == "Title")
    end = next(i for i in range(start + 1, len(paragraphs)) if p_or_none(paragraphs, i).text.strip() == "3. Command Console Reference" and p_or_none(paragraphs, i).style.name == "Title")
    for paragraph in paragraphs[start:end]:
        for run in paragraph.runs:
            run.text = run.text.replace("ZX81Emulator.exe", "ZX81EmulatorC.exe")


def p_or_none(paragraphs, index):
    return paragraphs[index]


def set_update_fields(document: Document) -> None:
    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("destination", type=Path)
    args = parser.parse_args()

    document = Document(args.source)
    replace_text_nodes(document)
    replace_heading_body(document, "/i", "Heading 1", add_language)
    replace_heading_body(document, "/r", "Heading 1", add_sound_startup)
    replace_heading_body(document, "/p", "Heading 1", add_port_or_peripherals)
    replace_heading_body(document, "/n", "Heading 1", add_ntsc)
    replace_heading_body(document, "/b", "Heading 1", add_border)
    replace_heading_body(document, "/w", "Heading 1", add_configuration)
    replace_heading_body(document, "/m", "Heading 1", add_machine)
    replace_family_sections(document)
    for name, command in STANDARD_OVERRIDES.items():
        replace_standard_command(document, name, command)
    for name, command in LOCAL_OVERRIDES.items():
        replace_standard_command(document, name, command, "Local console only.")
    replace_narrative(document)
    correct_console_examples(document)
    normalize_bullets(document)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        document.styles[style_name].paragraph_format.keep_with_next = True
    document.styles["Heading 2"].paragraph_format.space_before = Pt(14)
    for paragraph in document.paragraphs:
        if paragraph.style.name in {"Title", "Heading 1", "Heading 2", "Heading 3"}:
            paragraph.paragraph_format.keep_with_next = True
        if paragraph.style.name == "Heading 2" or (paragraph.style.name == "Heading 1" and paragraph.text.strip().startswith("/")):
            paragraph.paragraph_format.space_before = Pt(14)

    set_update_fields(document)
    document.core_properties.title = "Sinclair ZX80/ZX81 Emulator User Guide"
    document.core_properties.subject = "ZX80 and ZX81 ROM variants: startup options and command console reference"
    args.destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.destination)


if __name__ == "__main__":
    main()
