from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


GLOBAL_REPLACEMENTS = (
    ("VIC20::Emulator", "C264::C264Emulator"),
    ("VIC20Emulator_UserGuide.docx", "C264Emulator_UserGuide.docx"),
    ("VIC20EmulatorC.exe", "C264EmulatorC.exe"),
    ("VIC20Emulator.exe", "C264Emulator.exe"),
    ("VIC20EmulatorC", "C264EmulatorC"),
    ("VIC20Emulator", "C264Emulator"),
    ("VIC20::", "C264::"),
    ("VIC20-Specific", "C264-Specific"),
    ("VIC20-specific", "C264-specific"),
    ("VIC20 builder", "C264 builder"),
    ("Commodore VIC-20", "Commodore 264 Series"),
    ("VIC-20", "C264 series"),
    ("6502", "7501"),
    ("vic20.log", "c264.log"),
    ("vic20-screen.png", "c264-screen.png"),
)


C264_COMMANDS = (
    {
        "name": "SCREENDUMP",
        "aliases": "CSCREENDUMP.",
        "meaning": "Returns the current TED text-screen memory in hexadecimal.",
        "syntax": "SCREENDUMP",
        "parameters": ["No parameters."],
        "examples": ["SCREENDUMP"],
        "result": "Returns the screen-memory snapshot selected by TED and visible to the active CPU.",
    },
    {
        "name": "ATTRIBUTEDUMP",
        "aliases": "CATTRIBUTEDUMP.",
        "meaning": "Returns the current TED attribute memory in hexadecimal.",
        "syntax": "ATTRIBUTEDUMP",
        "parameters": ["No parameters."],
        "examples": ["ATTRIBUTEDUMP"],
        "result": "Returns the attribute-memory snapshot used by the active TED display mode.",
    },
    {
        "name": "BITMAPDUMP",
        "aliases": "CBITMAPDUMP.",
        "meaning": "Returns the current TED bitmap memory in hexadecimal.",
        "syntax": "BITMAPDUMP",
        "parameters": ["No parameters."],
        "examples": ["BITMAPDUMP"],
        "result": "Returns the bitmap-memory snapshot selected by TED and visible to the active CPU.",
    },
    {
        "name": "CHARSDRAW",
        "aliases": "CCHARSDRAW.",
        "meaning": "Renders textual drawings of all or selected character glyphs.",
        "syntax": "CHARSDRAW [0..255 ...]",
        "parameters": [
            "Character codes: optional values from 0 through 255; invalid values are ignored and duplicates are coalesced. With none, all characters are rendered."
        ],
        "examples": ["CHARSDRAW", "CHARSDRAW 1 2 65"],
        "result": "Returns textual drawings generated from the character data selected by TED.",
    },
)


INHERITED_RESULTS = {
    "VICII": (
        "Requests VIC-II state.",
        "Accepted by the inherited builder, but C264-series machines have no VIC-II, so no chip data is returned.",
    ),
    "VICIIEVENTS": (
        "Enables or disables VIC-II event visualization.",
        "Accepted by the inherited builder, but C264-series machines have no VIC-II, so the command has no effect.",
    ),
    "SID": (
        "Requests SID state.",
        "Accepted by the inherited builder, but C264-series sound is provided by TED rather than a SID chip, so no SID data is returned.",
    ),
    "SIDW": (
        "Changes the active SID sound wrapper.",
        "Accepted by the inherited builder, but C264-series machines have no SID chip; the TED sound implementation is unchanged.",
    ),
    "VICI": (
        "Requests VIC-I state.",
        "Accepted by the inherited builder, but C264-series machines have no VIC-I, so no chip data is returned.",
    ),
    "TED": (
        "Shows TED video, raster, interrupt and sound state.",
        "Returns the TED state used by the selected C16, C116 or Plus/4 model.",
    ),
    "TEDEVENTS": (
        "Enables or disables TED event visualization.",
        "Changes TED event visualization for the selected C264-series model.",
    ),
    "VIA": (
        "Requests generic Commodore VIA state.",
        "Accepted by the inherited builder, but the implemented C264-series machines have no VIA, so no VIA data is returned.",
    ),
    "CIA": (
        "Requests generic Commodore CIA state.",
        "Accepted by the inherited builder, but C264-series machines have no CIA. The Plus/4 ACIA is a different chip and is not exposed by this command.",
    ),
}


GRID_COMMANDS = {
    "GRIDON": {
        "aliases": "CGRIDON.",
        "meaning": "Enables the C264-specific screen grid and optionally TED raster-interrupt markers.",
        "syntax": "GRIDON COLOR [ON]",
        "parameters": [
            "COLOR: mandatory numeric grid color.",
            "ON: optional second token; its presence enables TED raster-interrupt-position markers.",
        ],
        "examples": ["GRIDON 14", "GRIDON 14 ON"],
        "result": "The C264-specific implementation shadows the standard GRIDON command. It enables the grid and, when a second token is supplied, TED raster-interrupt markers.",
    },
    "GRIDOFF": {
        "aliases": "CGRIDOFF.",
        "meaning": "Disables the C264-specific screen grid and TED raster-interrupt markers.",
        "syntax": "GRIDOFF",
        "parameters": ["No parameters."],
        "examples": ["GRIDOFF"],
        "result": "The C264-specific implementation shadows the standard GRIDOFF command and disables both the grid and TED raster-interrupt markers.",
    },
}


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_before(anchor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addprevious(element)
    paragraph = Paragraph(element, anchor._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def insert_labeled(anchor: Paragraph, label: str, value: str) -> Paragraph:
    paragraph = insert_before(anchor, style="Normal")
    paragraph.add_run(label).bold = True
    paragraph.add_run(value)
    return paragraph


def add_reference_body(anchor: Paragraph, command: dict[str, object], availability: str) -> None:
    insert_labeled(anchor, "Availability: ", availability)
    insert_labeled(anchor, "Accepted aliases: ", str(command["aliases"]))
    insert_labeled(anchor, "Meaning: ", str(command["meaning"]))
    insert_labeled(anchor, "General syntax: ", "")
    insert_before(anchor, str(command["syntax"]), "Code")
    insert_labeled(anchor, "Parameters: ", "")
    for parameter in command["parameters"]:
        insert_before(anchor, f"• {parameter}", "List Paragraph")
    insert_labeled(anchor, "Examples: ", "")
    for example in command["examples"]:
        insert_before(anchor, str(example), "Code")
    insert_labeled(anchor, "Result and consequences: ", str(command["result"]))


def add_command_block(anchor: Paragraph, command: dict[str, object]) -> None:
    heading = insert_before(anchor, str(command["name"]), "Heading 2")
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.keep_with_next = True
    add_reference_body(anchor, command, "Local console and remote /o channel.")


def replace_all_text_nodes(document: Document) -> None:
    for node in document.part.element.xpath(".//w:t"):
        if node.text is None:
            continue
        value = node.text
        for old, new in GLOBAL_REPLACEMENTS:
            value = value.replace(old, new)
        node.text = value


def heading_indices(document: Document, text: str, style: str) -> list[int]:
    return [
        i for i, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() == text and paragraph.style.name == style
    ]


def replace_parameter_blocks(document: Document, name: str, end_name: str, builder) -> None:
    occurrence = 0
    while True:
        paragraphs = document.paragraphs
        starts = heading_indices(document, name, "Heading 1")
        if occurrence >= len(starts):
            break
        start = starts[occurrence]
        end = next(
            i for i in range(start + 1, len(paragraphs))
            if paragraphs[i].text.strip() == end_name and paragraphs[i].style.name == "Heading 1"
        )
        heading = paragraphs[start]
        anchor = paragraphs[end]
        for paragraph in paragraphs[start + 1:end]:
            remove_paragraph(paragraph)
        heading.paragraph_format.space_before = Pt(14)
        heading.paragraph_format.keep_with_next = True
        builder(anchor, occurrence)
        occurrence += 1


def add_language_body(anchor: Paragraph, occurrence: int) -> None:
    executable = "C264Emulator.exe" if occurrence == 0 else "C264EmulatorC.exe"
    insert_labeled(anchor, "Meaning: ", "Selects the C264-series language and KERNAL-ROM variant.")
    insert_labeled(anchor, "General form: ", "")
    insert_before(anchor, "/iLANG", "Code")
    for value in (
        "ENG: English (default).",
        "FRA: French.",
        "SWE: Swedish.",
        "HUN: Hungarian.",
    ):
        insert_before(anchor, f"• {value}", "List Paragraph")
    insert_labeled(anchor, "Examples: ", "")
    insert_before(anchor, f"{executable} /iENG", "Code")
    insert_before(anchor, f"{executable} /iFRA", "Code")
    insert_labeled(anchor, "Consequence: ", "The selected language is used when loading the KERNAL ROM for the selected C16, C116 or Plus/4 model.")


def add_w_body(anchor: Paragraph, occurrence: int) -> None:
    executable = "C264Emulator.exe" if occurrence == 0 else "C264EmulatorC.exe"
    insert_labeled(anchor, "Meaning: ", "Selects the startup RAM configuration for the requested C264-series model.")
    insert_labeled(anchor, "General form: ", "")
    insert_before(anchor, "/w[CONF]", "Code")
    for value in (
        "C16 or C116, 0 or omitted: 16 KiB configuration (default).",
        "C16 or C116, 1: 32 KiB configuration.",
        "C16 or C116, 2 or any greater numeric value: 64 KiB configuration; values above 2 are clamped to 2.",
        "Plus/4: whenever /w is present, the emulator forces configuration 2 (64 KiB with the internal 3+1 ROM available), regardless of the numeric value supplied.",
    ):
        insert_before(anchor, f"• {value}", "List Paragraph")
    insert_labeled(anchor, "Examples: ", "")
    for example in (
        f"{executable} /mC16 /w0",
        f"{executable} /mC16 /w1",
        f"{executable} /mC116 /w2",
        f"{executable} /mCP4 /w2",
    ):
        insert_before(anchor, example, "Code")
    insert_labeled(anchor, "Consequence: ", "The selected RAM map is installed before machine initialization. For reliable Plus/4 startup, specify /w2: when /w is omitted, the current configuration selector returns 0 even though Plus/4 memory accepts only mode 2 and logs other modes as invalid.")


def insert_machine_blocks(document: Document) -> None:
    chapter_titles = (
        "2. Startup Parameters: Version With a Local Console",
        "3. Command Console Reference",
    )
    for occurrence, next_title in enumerate(chapter_titles):
        paragraphs = document.paragraphs
        start_search = 0 if occurrence == 0 else next(
            i for i, p in enumerate(paragraphs)
            if p.text.strip() == "2. Startup Parameters: Version With a Local Console" and p.style.name == "Title"
        )
        anchor_index = next(
            i for i in range(start_search, len(paragraphs))
            if paragraphs[i].text.strip() == next_title and paragraphs[i].style.name == "Title"
        )
        anchor = paragraphs[anchor_index]
        heading = insert_before(anchor, "/m", "Heading 1")
        heading.paragraph_format.space_before = Pt(14)
        heading.paragraph_format.keep_with_next = True
        executable = "C264Emulator.exe" if occurrence == 0 else "C264EmulatorC.exe"
        insert_labeled(anchor, "Meaning: ", "Selects the member of the Commodore 264 series to emulate.")
        insert_labeled(anchor, "General form: ", "")
        insert_before(anchor, "/mMACHINE", "Code")
        for value in (
            "C16: Commodore 16 implementation (default).",
            "C116: accepted as Commodore 116, but the current factory constructs the same Commodore16_116 implementation and C16/116 screen as C16.",
            "CP4: Commodore Plus/4 with 64 KiB RAM, ACIA, a second 6529B interface and the internal 3+1 ROM images.",
            "Any other value: falls back to C16.",
        ):
            insert_before(anchor, f"• {value}", "List Paragraph")
        insert_labeled(anchor, "Examples: ", "")
        for example in (
            f"{executable} /mC16",
            f"{executable} /mC116 /w2",
            f"{executable} /mCP4 /w2",
        ):
            insert_before(anchor, example, "Code")
        insert_labeled(anchor, "Consequence: ", "The model choice determines the computer class, ROM set, memory implementation, screen label and model-specific chips created before initialization.")


def replace_command_body(document: Document, name: str, command: dict[str, object]) -> None:
    paragraphs = document.paragraphs
    start = next(
        i for i, p in enumerate(paragraphs)
        if p.text.strip() == name and p.style.name == "Heading 2"
    )
    end = next(
        i for i in range(start + 1, len(paragraphs))
        if paragraphs[i].style.name in {"Heading 2", "Heading 1", "Title"}
    )
    anchor = paragraphs[end]
    for paragraph in paragraphs[start + 1:end]:
        remove_paragraph(paragraph)
    add_reference_body(anchor, command, "Local console and remote /o channel.")


def tailor_inherited_commands(document: Document) -> None:
    paragraphs = document.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "3.2 Inherited Commodore Commands" and p.style.name == "Heading 1")
    end = next(i for i in range(start + 1, len(paragraphs)) if paragraphs[i].text.strip().startswith("3.3 ") and paragraphs[i].style.name == "Heading 1")
    indices = [i for i in range(start + 1, end) if paragraphs[i].style.name == "Heading 2"] + [end]
    for position in range(len(indices) - 1):
        command = paragraphs[indices[position]].text.strip()
        if command not in INHERITED_RESULTS:
            continue
        meaning, result = INHERITED_RESULTS[command]
        for paragraph in paragraphs[indices[position] + 1:indices[position + 1]]:
            if paragraph.text.startswith("Meaning:"):
                paragraph.clear()
                paragraph.add_run("Meaning: ").bold = True
                paragraph.add_run(meaning)
            elif paragraph.text.startswith("Result and consequences:"):
                paragraph.clear()
                paragraph.add_run("Result and consequences: ").bold = True
                paragraph.add_run(result)


def replace_specific_section(document: Document) -> None:
    paragraphs = document.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith("3.3 ") and p.style.name == "Heading 1")
    end = next(i for i in range(start + 1, len(paragraphs)) if paragraphs[i].text.strip() == "3.4 Local-Console-Only Commands" and paragraphs[i].style.name == "Heading 1")
    anchor = paragraphs[end]
    for paragraph in paragraphs[start:end]:
        remove_paragraph(paragraph)
    section = insert_before(anchor, "3.3 C264-Specific Commands", "Heading 1")
    section.paragraph_format.keep_with_next = True
    for command in C264_COMMANDS:
        add_command_block(anchor, command)


def replace_narrative(document: Document) -> None:
    replacements = {
        "This guide describes the startup interface": (
            "This guide describes the startup and command interfaces implemented by the two Commodore 264-series executables in EMULATORS: C264Emulator (without a local command console) and C264EmulatorC (with a local command console). It covers the Commodore 16, Commodore 116 and Commodore Plus/4 models and is based on both executable entry points, the C264::C264Emulator inheritance chain, LocalConsole, and the complete C264 -> COMMODORE -> StandardCommandBuilder responsibility chain."
        ),
        "Executable: C264Emulator.exe.": (
            "Executable: C264Emulator.exe. It accepts the framework and C264-specific options below, plus /o and the communication-port option. Option letters are case-sensitive because the command-line parser stores the character exactly as supplied."
        ),
        "Executable: C264EmulatorC.exe.": (
            "Executable: C264EmulatorC.exe. Its main program delegates the complete command line to C264::C264Emulator, so all inherited framework options and C264-specific options are active even though the short main-level help banner only advertises /h. It has no /o listener option; commands are entered in its local console."
        ),
        "The local console checks its eleven": (
            "The local console checks its eleven emulator-level commands first. All remaining text is delegated through C264::CommandBuilder, then COMMODORE::CommandBuilder, and finally MCHEmul::StandardCommandBuilder. The remote /o channel in C264Emulator.exe uses only that three-builder chain. The availability line in each entry makes this boundary explicit."
        ),
        "This file is a maintained repository artifact.": (
            "This file is a maintained repository artifact. Every change that adds, removes, renames or changes a C264 startup option, a command-builder registration, a LocalConsole-only command, command syntax, parameter meaning, command consequences or user-visible formatted output must include a review and, when applicable, an update of docs/C264Data/C264Emulator_UserGuide.docx. The corresponding canonical .fmt source must be updated whenever the command's InfoStructure contract changes."
        ),
        "Audit the complete chain, not only the edited class:": (
            "Audit the complete chain, not only the edited class: the C264Emulator entry point and C264::C264Emulator hierarchy for startup options; LocalConsole for local-only commands; and C264::CommandBuilder -> COMMODORE::CommandBuilder -> MCHEmul::StandardCommandBuilder for builder commands."
        ),
    }
    for paragraph in document.paragraphs:
        for prefix, value in replacements.items():
            if paragraph.text.startswith(prefix):
                paragraph.clear()
                paragraph.add_run(value)
                break


def set_update_fields(document: Document) -> None:
    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("destination", type=Path)
    args = parser.parse_args()

    document = Document(args.source)
    replace_all_text_nodes(document)
    replace_parameter_blocks(document, "/i", "/r", add_language_body)
    replace_parameter_blocks(document, "/w", "/b", add_w_body)
    insert_machine_blocks(document)
    for name, command in GRID_COMMANDS.items():
        replace_command_body(document, name, command)
    tailor_inherited_commands(document)
    replace_specific_section(document)
    replace_narrative(document)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        document.styles[style_name].paragraph_format.keep_with_next = True
    document.styles["Heading 2"].paragraph_format.space_before = Pt(14)
    for paragraph in document.paragraphs:
        if paragraph.style.name in {"Title", "Heading 1", "Heading 2", "Heading 3"}:
            paragraph.paragraph_format.keep_with_next = True
        if paragraph.style.name == "Heading 2" or (paragraph.style.name == "Heading 1" and paragraph.text.strip().startswith("/")):
            paragraph.paragraph_format.space_before = Pt(14)

    set_update_fields(document)
    document.core_properties.title = "Commodore 264 Series Emulator User Guide"
    document.core_properties.subject = "C16, C116 and Plus/4 startup options and command console reference"
    args.destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.destination)


if __name__ == "__main__":
    main()
