from __future__ import annotations

import argparse
import importlib.util
from pathlib import Path

from docx import Document
from docx.shared import Pt


BASE_PATH = Path(r"C:\Workspaces\EMULATORS\outputs\zx81-user-guide-work\build_zx81_guide.py")
SPEC = importlib.util.spec_from_file_location("zx81_guide_support", BASE_PATH)
ZX = importlib.util.module_from_spec(SPEC)
assert SPEC.loader is not None
SPEC.loader.exec_module(ZX)
SUPPORT = ZX.SUPPORT


REPLACEMENTS = (
    ("Sinclair ZX80/ZX81 Emulator User Guide", "Sinclair ZX Spectrum Emulator User Guide"),
    ("ZX80 and ZX81 ROM Variants - Startup and Console Reference", "ZX Spectrum 16K/48K - Startup and Console Reference"),
    ("ZX80 and ZX81 ROM Variants", "ZX Spectrum 16K/48K"),
    ("ZX81Emulator_UserGuide.docx", "ZXSpectrumEmulator_UserGuide.docx"),
    ("ZX81EmulatorC.exe", "ZXSpectrumEmulatorC.exe"),
    ("ZX81Emulator.exe", "ZXSpectrumEmulator.exe"),
    ("ZX81EmulatorC", "ZXSpectrumEmulatorC"),
    ("ZX81Emulator", "ZXSpectrumEmulator"),
    ("ZX81::ZX81Emulator", "ZXSPECTRUM::ZXSpectrumEmulator"),
    ("ZX81::CommandBuilder", "ZXSPECTRUM::CommandBuilder"),
    ("ZX80/ZX81-specific", "ZX Spectrum-specific"),
    ("ZX80/ZX81-Specific", "ZX Spectrum-Specific"),
    ("ZX80/ZX81 builder", "ZX Spectrum builder"),
    ("Sinclair ZX80/ZX81", "Sinclair ZX Spectrum"),
    ("ZX80/ZX81", "ZX Spectrum"),
    ("ZX80 or ZX81", "ZX Spectrum"),
    ("ZX80 and ZX81", "ZX Spectrum"),
    ("zx81.log", "zxspectrum.log"),
    ("zx81-screen.png", "zxspectrum-screen.png"),
    ("game.p", "game.tap"),
)


SINCLAIR_COMMANDS = (
    {
        "name": "SYSVARS", "aliases": "CSYSVARS.",
        "meaning": "Returns every system variable defined for the selected ZX Spectrum ROM profile.",
        "syntax": "SYSVARS", "parameters": ["No parameters."], "examples": ["SYSVARS"],
        "result": "Returns the current values using the system-variable definitions loaded for the selected ROM language.",
    },
    {
        "name": "SYSVAR", "aliases": "CSYSVAR.",
        "meaning": "Returns one named ZX Spectrum system variable.",
        "syntax": "SYSVAR VNAME",
        "parameters": ["VNAME: mandatory variable name from the active ZX Spectrum system-variable table."],
        "examples": ["SYSVAR FRAMES", "SYSVAR RAMTOP"],
        "result": "Returns the named variable and its current memory value; an unknown name returns an error entry.",
    },
)


SPECTRUM_COMMANDS = (
    {
        "name": "ULA", "aliases": "CULA.",
        "meaning": "Returns the current ZX Spectrum ULA and video-generation state.",
        "syntax": "ULA", "parameters": ["No parameters."], "examples": ["ULA"],
        "result": "Returns the active ULA state, including raster and display-generation information for PAL or NTSC operation.",
    },
    {
        "name": "ULAEVENTS", "aliases": "CULAEVENTS.",
        "meaning": "Enables or disables visualization of significant CPU/ULA events.",
        "syntax": "ULAEVENTS ON|OFF",
        "parameters": ["ON: enables event visualization.", "OFF: disables it. The implementation accepts any single token; only the exact value ON enables the feature."],
        "examples": ["ULAEVENTS ON", "ULAEVENTS OFF"],
        "result": "Changes the event-overlay flag used by the active ZX Spectrum ULA.",
    },
    {
        "name": "SCREENDUMP", "aliases": "CSCREENDUMP.",
        "meaning": "Returns a snapshot of ZX Spectrum bitmap memory.",
        "syntax": "SCREENDUMP", "parameters": ["No parameters."], "examples": ["SCREENDUMP"],
        "result": "Returns the current 6,144-byte bitmap area selected by the machine's screen-memory snapshot routine.",
    },
    {
        "name": "COLORDUMP", "aliases": "CCOLORDUMP.",
        "meaning": "Returns a snapshot of ZX Spectrum attribute memory.",
        "syntax": "COLORDUMP", "parameters": ["No parameters."], "examples": ["COLORDUMP"],
        "result": "Returns the current 768-byte colour-attribute area selected by the machine's colour-memory snapshot routine.",
    },
    {
        "name": "CHARSDRAW", "aliases": "CCHARSDRAW.",
        "meaning": "Returns textual drawings of all or selected ZX Spectrum character glyphs.",
        "syntax": "CHARSDRAW [0..95 ...]",
        "parameters": ["Character indexes: optional values from 0 through 95; out-of-range values are ignored and duplicates are coalesced. With none, all 96 glyphs are rendered."],
        "examples": ["CHARSDRAW", "CHARSDRAW 0 31 64 95"],
        "result": "Returns textual 8x8 glyph drawings from the standard English 48K ROM character data used by the screen object.",
    },
)


STANDARD_OVERRIDES = {
    "SOUNDON": {
        "name": "SOUNDON", "aliases": "CSOUNDON.",
        "meaning": "Enables ZX Spectrum ULA sound output.", "syntax": "SOUNDON",
        "parameters": ["No parameters."], "examples": ["SOUNDON"],
        "result": "Enables the sound system attached to the ULA sound function.",
    },
    "SOUNDOFF": {
        "name": "SOUNDOFF", "aliases": "CSOUNDOFF.",
        "meaning": "Disables ZX Spectrum ULA sound output.", "syntax": "SOUNDOFF",
        "parameters": ["No parameters."], "examples": ["SOUNDOFF"],
        "result": "Disables the sound system; emulation continues without audible ULA output.",
    },
}


LOCAL_OVERRIDES = {
    "POSSIBLEPERS": {
        "name": "POSSIBLEPERS", "aliases": "None.",
        "meaning": "Lists the peripheral IDs and construction attributes advertised by the ZX Spectrum I/O peripheral builder.",
        "syntax": "POSSIBLEPERS", "parameters": ["No parameters."], "examples": ["POSSIBLEPERS"],
        "result": "Lists IDs 10 (typewriter), 100 (cassette), 101 (cassette injection) and 102 (thermal printer). Cartridge ID 201 is accepted by the builder but is not advertised.",
    },
    "CONNECTPER": {
        "name": "CONNECTPER", "aliases": "None.",
        "meaning": "Builds and connects a peripheral to the first compatible ZX Spectrum device.",
        "syntax": "CONNECTPER ID [ATTRS...]",
        "parameters": [
            "ID: peripheral-emulation ID (10, 100, 101, 102 or the additionally supported cartridge ID 201).",
            "ATTRS: optional construction attributes. For printer 102, use f:FILENAME and p:THERMAL or p:THERMAL-PS.",
        ],
        "examples": ["CONNECTPER 100", "CONNECTPER 102 f:printer.ps p:THERMAL-PS"],
        "result": "Connection succeeds only when the ID, attributes and a compatible target device are available; otherwise the answer reports an error.",
    },
    "LOADPERDATA": {
        "name": "LOADPERDATA", "aliases": "None.",
        "meaning": "Loads media or program data into a connected ZX Spectrum peripheral.",
        "syntax": "LOADPERDATA ID FILE [DEBUGFILE]",
        "parameters": [
            "ID: connected peripheral ID.",
            "FILE: supported data file; for example sampled TZX data for ID 100, or TAP/TAP-like TZX data for injection ID 101.",
            "DEBUGFILE: optional file that starts full-range format-related debugging after a successful load.",
        ],
        "examples": ["LOADPERDATA 100 tape.tzx", "LOADPERDATA 101 game.tap load-trace.log"],
        "result": "The operation fails if the peripheral rejects the file type. Injection accepts TAP and only TZX files whose blocks are all TAP-compatible.",
    },
    "EMPTYPERDATA": {
        "name": "EMPTYPERDATA", "aliases": "None.",
        "meaning": "Creates and attaches empty media data when the peripheral implements that operation.",
        "syntax": "EMPTYPERDATA ID FILE",
        "parameters": ["ID: connected peripheral ID.", "FILE: destination whose extension selects a writable format."],
        "examples": ["EMPTYPERDATA 100 blank.tzx"],
        "result": "The Sinclair cassette can create an empty TZX container; unsupported peripherals or formats return an error.",
    },
    "SAVEPERDATA": {
        "name": "SAVEPERDATA", "aliases": "None.",
        "meaning": "Saves a peripheral's current media data.",
        "syntax": "SAVEPERDATA ID [FILE]",
        "parameters": ["ID: connected peripheral ID.", "FILE: optional output path; it may be omitted when a prior load/create operation established the filename."],
        "examples": ["SAVEPERDATA 100", "SAVEPERDATA 100 copy.tzx"],
        "result": "Saves current cassette data when the peripheral can retrieve it and the selected format is writable. Injection ID 101 cannot retrieve data.",
    },
}


def replace_text_nodes(document: Document) -> None:
    for node in document.part.element.xpath(".//w:t"):
        if node.text is None:
            continue
        value = node.text
        for old, new in REPLACEMENTS:
            value = value.replace(old, new)
        node.text = value


def executable(occurrence: int) -> str:
    return "ZXSpectrumEmulator.exe" if occurrence == 0 else "ZXSpectrumEmulatorC.exe"


def add_language(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Selects the ZX Spectrum ROM language.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/iLANG", "Code")
    for text in ("ENG: English (default).", "ESP: Spanish.", "NOR: Nordic.", "ARA: Arabic."):
        ZX.bullet(anchor, text)
    ZX.bullet(anchor, "Language identifiers are case-sensitive; use the uppercase forms above. Any other value loads the English ROM.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /iESP", "Code")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /iARA", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "The selected 16 KiB 48K-family ROM image is loaded before initialization; the screen's glyph reference remains sourced from the English ROM file.")


def add_sound_startup(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Selects the initial ULA sound-output state.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/r[YES|NO]", "Code")
    ZX.bullet(anchor, "YES or an empty value: sound enabled.")
    ZX.bullet(anchor, "NO (and any other non-YES non-empty value): sound disabled.")
    ZX.bullet(anchor, "Omitted: sound enabled.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /r", "Code")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /rNO", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "The ZX Spectrum sound device attached to the ULA sound function starts enabled or disabled as requested.")


def add_ntsc(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Selects NTSC timing and screen geometry instead of PAL.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/n", "Code")
    ZX.bullet(anchor, "No value: the presence of the option selects NTSC; omission selects PAL.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /n", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "The factory creates the NTSC ULA, NTSC screen and 3.5275 MHz clock. Without /n it creates the PAL ULA, PAL screen and 3.5 MHz clock.")


def add_border(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Draws the diagnostic character-cell grid over the active screen and selects its startup colour.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/b[COLOR]", "Code")
    ZX.bullet(anchor, "COLOR omitted or 0: colour 0 (default).")
    ZX.bullet(anchor, "1 through 3: the requested startup grid colour.")
    ZX.bullet(anchor, "Any value greater than 3: clamped to 3 by the startup parser.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /b", "Code")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /b3", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "The screen starts with its diagnostic 8x8 grid enabled. The later GRIDON console command has a wider 0-15 renderer range.")


def add_configuration(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Selects the installed RAM configuration.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/wCONF", "Code")
    ZX.bullet(anchor, "0 or omitted: 16 KiB RAM configuration (default).")
    ZX.bullet(anchor, "1: 48 KiB RAM configuration.")
    ZX.bullet(anchor, "2 or any value greater than 1: unsupported. The current range check can pass or clamp it to the invalid internal value 2, which may assert in debug builds and must not be used.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /w0", "Code")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /w1", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "Configuration 0 maps 16 KiB RAM with unmapped upper memory; configuration 1 maps the full 48 KiB RAM area before ROM startup.")


def add_machine(anchor, occurrence: int) -> None:
    SUPPORT.insert_labeled(anchor, "Meaning: ", "Selects the ZX Spectrum machine family implemented by this emulator.")
    SUPPORT.insert_labeled(anchor, "General form: ", "")
    SUPPORT.insert_before(anchor, "/mMACHINE", "Code")
    ZX.bullet(anchor, "ZXSTD: standard 16K/48K ZX Spectrum; this is also the default.")
    ZX.bullet(anchor, "Any other value: silently falls back to ZXSTD in the current implementation.")
    SUPPORT.insert_labeled(anchor, "Examples: ", "")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /mZXSTD", "Code")
    SUPPORT.insert_before(anchor, f"{executable(occurrence)} /mZXSTD /w1 /iESP", "Code")
    SUPPORT.insert_labeled(anchor, "Consequence: ", "Only the original standard ZX Spectrum type is created. This selector does not provide 128K, +2, +2A or +3 models; /w, /i and /n independently select RAM, ROM language and video standard.")


def add_port_or_peripherals(anchor, occurrence: int) -> None:
    if occurrence == 0:
        SUPPORT.insert_labeled(anchor, "Meaning: ", "Selects the TCP listening port used by the non-console executable's remote command channel.")
        SUPPORT.insert_labeled(anchor, "General form: ", "")
        SUPPORT.insert_before(anchor, "/pPORTNUMBER", "Code")
        ZX.bullet(anchor, "0 through 65535: requested listening port; values above 1000 are recommended.")
        ZX.bullet(anchor, "Omitted: port 60000 is used when /o enables the channel.")
        ZX.bullet(anchor, "Collision: the inherited emulator also interprets /p as a comma-separated peripheral-ID list. A supplied port number is therefore also attempted as a peripheral ID during initialization.")
        SUPPORT.insert_labeled(anchor, "Examples: ", "")
        SUPPORT.insert_before(anchor, "ZXSpectrumEmulator.exe /o", "Code")
        SUPPORT.insert_before(anchor, "ZXSpectrumEmulator.exe /o /p61000", "Code")
        SUPPORT.insert_labeled(anchor, "Consequence: ", "The remote server listens on the selected port. When /p is present, the same value is also attempted as a peripheral ID and normally produces a log warning.")
    else:
        SUPPORT.insert_labeled(anchor, "Meaning: ", "Connects one or more peripherals during startup.")
        SUPPORT.insert_labeled(anchor, "General form: ", "")
        SUPPORT.insert_before(anchor, "/pPER1[,PER2...]", "Code")
        for value in (
            "10: typewriter keyboard input.",
            "100: Sinclair cassette signal simulation.",
            "101: ZX Spectrum TAP/TZX cassette-data injection.",
            "102: ZX Printer thermal-printer simulation.",
            "201: cartridge implementation supported by the builder, although it is not listed by POSSIBLEPERS.",
        ):
            ZX.bullet(anchor, value)
        SUPPORT.insert_labeled(anchor, "Examples: ", "")
        SUPPORT.insert_before(anchor, "ZXSpectrumEmulatorC.exe /p10", "Code")
        SUPPORT.insert_before(anchor, "ZXSpectrumEmulatorC.exe /p100,102", "Code")
        SUPPORT.insert_labeled(anchor, "Consequence: ", "Each requested peripheral is built and connected to the first compatible device during initialization. Unsupported IDs are logged and initialization continues.")


def replace_family_sections(document: Document) -> None:
    paragraphs = document.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith("3.2 ") and p.style.name == "Heading 1")
    end = next(i for i in range(start + 1, len(paragraphs)) if paragraphs[i].text.strip() == "3.4 Local-Console-Only Commands" and paragraphs[i].style.name == "Heading 1")
    anchor = paragraphs[end]
    for paragraph in paragraphs[start:end]:
        SUPPORT.remove_paragraph(paragraph)
    heading = SUPPORT.insert_before(anchor, "3.2 Inherited Sinclair Commands", "Heading 1")
    heading.paragraph_format.keep_with_next = True
    for command in SINCLAIR_COMMANDS:
        ZX.add_command_block(anchor, command)
    heading = SUPPORT.insert_before(anchor, "3.3 ZX Spectrum-Specific Commands", "Heading 1")
    heading.paragraph_format.keep_with_next = True
    for command in SPECTRUM_COMMANDS:
        ZX.add_command_block(anchor, command)


def replace_narrative(document: Document) -> None:
    replacements = {
        "This guide describes the startup and command interfaces": "This guide describes the startup and command interfaces implemented by the two Sinclair ZX Spectrum executables in EMULATORS: ZXSpectrumEmulator (without a local command console) and ZXSpectrumEmulatorC (with a local command console). It covers the standard 16K/48K machine, PAL and NTSC video, four ROM languages, and is based on both entry points, the ZXSPECTRUM::ZXSpectrumEmulator inheritance chain, LocalConsole, and the complete ZXSPECTRUM -> SINCLAIR -> StandardCommandBuilder responsibility chain.",
        "Executable: ZXSpectrumEmulator.exe.": "Executable: ZXSpectrumEmulator.exe. It accepts the framework and ZX Spectrum-specific options below, plus /o and the communication-port option. Option letters and language values are case-sensitive because the command-line parser preserves them exactly.",
        "Executable: ZXSpectrumEmulatorC.exe.": "Executable: ZXSpectrumEmulatorC.exe. Its main program delegates the complete command line to ZXSPECTRUM::ZXSpectrumEmulator, so all inherited framework options and ZX Spectrum-specific options are active even though the short main-level banner only advertises /h. It has no /o listener; commands are entered in its local console.",
        "The local console checks its eleven": "The local console checks its eleven emulator-level commands first. All remaining text is delegated through ZXSPECTRUM::CommandBuilder, then SINCLAIR::CommandBuilder, and finally MCHEmul::StandardCommandBuilder. The remote /o channel in ZXSpectrumEmulator.exe uses only that three-builder chain. The availability line in each entry makes this boundary explicit.",
        "This file is a maintained repository artifact.": "This file is a maintained repository artifact. Every change that adds, removes, renames or changes a ZX Spectrum startup option, command-builder registration, LocalConsole-only command, syntax, parameter meaning, consequence or visible formatted output must include a review and, when applicable, an update of docs/ZXSpectrumData/ZXSpectrumEmulator_UserGuide.docx. Update the corresponding canonical .fmt source whenever the command's InfoStructure contract changes.",
        "Audit the complete chain, not only the edited class:": "Audit the complete chain, not only the edited class: both ZXSpectrumEmulator entry points and the ZXSPECTRUM::ZXSpectrumEmulator hierarchy for startup options; LocalConsole for local-only commands; and ZXSPECTRUM::CommandBuilder -> SINCLAIR::CommandBuilder -> MCHEmul::StandardCommandBuilder for builder commands.",
    }
    for paragraph in document.paragraphs:
        for prefix, value in replacements.items():
            if paragraph.text.startswith(prefix):
                paragraph.clear()
                paragraph.add_run(value)
                break


def correct_console_examples(document: Document) -> None:
    paragraphs = document.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "2. Startup Parameters: Version With a Local Console" and p.style.name == "Title")
    end = next(i for i in range(start + 1, len(paragraphs)) if paragraphs[i].text.strip() == "3. Command Console Reference" and paragraphs[i].style.name == "Title")
    for paragraph in paragraphs[start:end]:
        for run in paragraph.runs:
            run.text = run.text.replace("ZXSpectrumEmulator.exe", "ZXSpectrumEmulatorC.exe")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("destination", type=Path)
    args = parser.parse_args()

    document = Document(args.source)
    replace_text_nodes(document)
    ZX.replace_heading_body(document, "/i", "Heading 1", add_language)
    ZX.replace_heading_body(document, "/r", "Heading 1", add_sound_startup)
    ZX.replace_heading_body(document, "/p", "Heading 1", add_port_or_peripherals)
    ZX.replace_heading_body(document, "/n", "Heading 1", add_ntsc)
    ZX.replace_heading_body(document, "/b", "Heading 1", add_border)
    ZX.replace_heading_body(document, "/w", "Heading 1", add_configuration)
    ZX.replace_heading_body(document, "/m", "Heading 1", add_machine)
    replace_family_sections(document)
    for name, command in STANDARD_OVERRIDES.items():
        ZX.replace_standard_command(document, name, command)
    for name, command in LOCAL_OVERRIDES.items():
        ZX.replace_standard_command(document, name, command, "Local console only.")
    replace_narrative(document)
    correct_console_examples(document)
    ZX.normalize_bullets(document)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        document.styles[style_name].paragraph_format.keep_with_next = True
    document.styles["Heading 2"].paragraph_format.space_before = Pt(14)
    for paragraph in document.paragraphs:
        if paragraph.style.name in {"Title", "Heading 1", "Heading 2", "Heading 3"}:
            paragraph.paragraph_format.keep_with_next = True
        if paragraph.style.name == "Heading 2" or (paragraph.style.name == "Heading 1" and paragraph.text.strip().startswith("/")):
            paragraph.paragraph_format.space_before = Pt(14)

    ZX.set_update_fields(document)
    document.core_properties.title = "Sinclair ZX Spectrum Emulator User Guide"
    document.core_properties.subject = "ZX Spectrum 16K/48K: startup options and command console reference"
    args.destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.destination)


if __name__ == "__main__":
    main()
