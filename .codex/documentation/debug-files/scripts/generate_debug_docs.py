#!/usr/bin/env python3
"""Generate the maintained EMULATORS deep-debug reference guides."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
from pathlib import Path
import subprocess
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


HERE = Path(__file__).resolve().parent
PACKAGE = HERE.parent
SOURCES = PACKAGE / "sources"
ROOT = PACKAGE.parents[2]
COMMON_PATH = SOURCES / "debug-common.yaml"
MACHINE_PATHS = [
    SOURCES / "c64.yaml",
    SOURCES / "vic20.yaml",
    SOURCES / "c264.yaml",
    SOURCES / "zx80-zx81.yaml",
    SOURCES / "zxspectrum.yaml",
    SOURCES / "msx.yaml",
]
SKILL = ROOT / ".codex" / "skills" / "emulators-documentation-authoring"
STARTER = SKILL / "scripts" / "start_document.py"
STATE_PATH = PACKAGE / "state" / "generated-state.json"


def load_yaml_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as stream:
        return json.load(stream)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_paragraph_before(anchor, text: str = "", style: str | None = None):
    paragraph = OxmlElement("w:p")
    anchor._element.addprevious(paragraph)
    from docx.text.paragraph import Paragraph
    result = Paragraph(paragraph, anchor._parent)
    if style:
        result.style = style
    if text:
        result.add_run(text)
    return result


def keep_with_next(paragraph) -> None:
    props = paragraph._p.get_or_add_pPr()
    if props.find(qn("w:keepNext")) is None:
        props.append(OxmlElement("w:keepNext"))


def keep_together(paragraph) -> None:
    props = paragraph._p.get_or_add_pPr()
    if props.find(qn("w:keepLines")) is None:
        props.append(OxmlElement("w:keepLines"))


def remove_page_break_before(element) -> None:
    """Remove inherited/direct page-break-before without changing other spacing."""
    props = element.get_or_add_pPr() if hasattr(element, "get_or_add_pPr") else element._p.get_or_add_pPr()
    node = props.find(qn("w:pageBreakBefore"))
    if node is not None:
        props.remove(node)


def shade_cell(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = props.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        props.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margin(cell, top=90, start=110, bottom=90, end=110) -> None:
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    props = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    props.append(header)


def add_table_before(doc, anchor, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False
    table._tbl.getparent().remove(table._tbl)
    anchor._element.addprevious(table._tbl)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        shade_cell(cell, "1F4E78")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8.5)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[index].paragraphs:
                paragraph.style = "Normal"
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
        if len(table.rows) % 2 == 1:
            for cell in cells:
                shade_cell(cell, "EAF2F8")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[index])
            set_cell_margin(cell)
    table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def add_heading(anchor, text: str, level: int = 1):
    paragraph = insert_paragraph_before(anchor, text, f"Heading {level}")
    remove_page_break_before(paragraph)
    paragraph.paragraph_format.space_before = Pt(14 if level >= 2 else 10)
    keep_with_next(paragraph)
    return paragraph


def add_title(anchor, text: str):
    paragraph = insert_paragraph_before(anchor, text, "Title")
    keep_with_next(paragraph)
    return paragraph


def add_body(anchor, text: str, bold_lead: str | None = None):
    paragraph = insert_paragraph_before(anchor, style="Normal")
    paragraph.paragraph_format.space_after = Pt(6)
    if bold_lead and text.startswith(bold_lead):
        paragraph.add_run(bold_lead).bold = True
        paragraph.add_run(text[len(bold_lead):])
    else:
        paragraph.add_run(text)
    return paragraph


def create_num_instance(doc, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    wanted = "decimal" if ordered else "bullet"
    abstract_id = None
    for abstract in numbering.findall(qn("w:abstractNum")):
        level = abstract.find(qn("w:lvl"))
        if level is None:
            continue
        num_fmt = level.find(qn("w:numFmt"))
        if num_fmt is not None and num_fmt.get(qn("w:val")) == wanted:
            abstract_id = int(abstract.get(qn("w:abstractNumId")))
            break
    if abstract_id is None:
        raise RuntimeError(f"The template does not contain {wanted} numbering")
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_num_id(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def add_bullets(doc, anchor, items, numbered: bool = False):
    num_id = create_num_instance(doc, numbered)
    for item in items:
        paragraph = insert_paragraph_before(anchor, str(item), "Normal")
        set_num_id(paragraph, num_id)
        paragraph.paragraph_format.space_after = Pt(3)


def add_code(anchor, lines):
    paragraph = insert_paragraph_before(anchor, style="Normal")
    paragraph.paragraph_format.left_indent = Cm(0.35)
    paragraph.paragraph_format.right_indent = Cm(0.35)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.0
    props = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    props.append(shading)
    for index, line in enumerate(lines):
        run = paragraph.add_run(("" if index == 0 else "\n") + line)
        run.font.name = "Consolas"
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Consolas")
        run.font.size = Pt(7.5)
    keep_together(paragraph)


def set_update_fields(doc) -> None:
    settings = doc.settings._element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def clear_update_fields(doc) -> None:
    settings = doc.settings._element
    node = settings.find(qn("w:updateFields"))
    if node is not None:
        settings.remove(node)


def normalize_word_field_results(doc) -> None:
    """Translate cached empty-list results produced by a Spanish Word install."""
    replacements = {
        "NO SE ENCUENTRAN ELEMENTOS DE TABLA DE ILUSTRACIONES.": "No table of illustrations entries were found.",
        "No se encuentran elementos de tabla de ilustraciones.": "No table of illustrations entries were found.",
    }
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)


def configure_styles(doc) -> None:
    for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        props = style.element.get_or_add_pPr()
        if props.find(qn("w:keepNext")) is None:
            props.append(OxmlElement("w:keepNext"))
        if name != "Title":
            remove_page_break_before(style.element)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)


def replace_text_in_container(container, old: str, new: str) -> None:
    """Replace a label while preserving field-code runs in headers and footers."""
    for paragraph in container.paragraphs:
        for run in paragraph.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_container(cell, old, new)


def build_document(common: dict, machine: dict) -> Path:
    output = ROOT / machine["output"]
    output.parent.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [sys.executable, str(STARTER), "--title", machine["title"], "--subtitle", machine["subtitle"],
         "--author", common["author"], "--year", str(common["year"]), "--out", str(output), "--force"],
        check=True,
    )
    doc = Document(output)
    configure_styles(doc)
    doc.core_properties.subject = "EMULATORS deep-debug file format"
    doc.core_properties.keywords = f"EMULATORS, DEBUG, {machine['id']}, trace, emulation"
    doc.core_properties.comments = "Generated from .codex/documentation/debug-files/sources and the emulator-specific source."

    prologue = next(p for p in doc.paragraphs if p.text.startswith("[Escriba"))
    prologue.text = (
        f"This guide defines the DEBUG format for {machine['title'].replace('DEBUG File', '').strip()}. "
        f"Its purpose is to make a trace readable and auditable directly against the code that produces it. "
        f"Audience: {common['audience']}"
    )
    prologue.paragraph_format.space_after = Pt(6)

    for paragraph in list(doc.paragraphs):
        if paragraph.text.startswith("[Título del bloque") or paragraph.text.startswith("[Desarrolle"):
            remove_paragraph(paragraph)

    label_translations = {
        "Índice": "Contents",
        "Prólogo": "Foreword",
        "Anexos": "Appendices",
        "Ilustraciones": "Illustrations",
        "Ecuaciones": "Equations",
    }
    for paragraph in doc.paragraphs:
        if paragraph.text in label_translations:
            paragraph.text = label_translations[paragraph.text]
    for section in doc.sections:
        replace_text_in_container(section.header, "Página", "Page")
        replace_text_in_container(section.footer, "Página", "Page")

    annex = next(p for p in doc.paragraphs if p.text == "Appendices")
    body_children = list(doc.element.body)
    prologue_pos = body_children.index(prologue._p)
    annex_pos = body_children.index(annex._p)
    for paragraph in list(doc.paragraphs):
        position = body_children.index(paragraph._p)
        if paragraph.text == "" and prologue_pos < position < annex_pos:
            remove_paragraph(paragraph)


    add_title(annex, "1. Purpose, scope, and models")
    add_body(annex, common["purpose"])
    add_heading(annex, "1.1 Scope of this guide")
    add_body(annex, f"Covered models: {machine['models']}")
    add_body(annex, f"Main processor: {machine['cpu']}")
    add_body(annex, "The guide describes the behavior implemented in the current repository revision. When a component is present but its producer contains a TODO, the guide explicitly states that it currently produces no output.")

    add_title(annex, "2. Activation and producer selection")
    add_heading(annex, "2.1 Console command")
    add_code(annex, [common["activation"]["command"], common["activation"]["off_command"], common["activation"]["interrupt_on"], common["activation"]["interrupt_off"]])
    add_table_before(doc, annex, ["Parameter", "Meaning"], common["activation"]["parameters"], [4.0, 11.7])
    add_heading(annex, "2.2 Conditions and special paths")
    add_bullets(doc, annex, common["activation"]["notes"])

    add_title(annex, "3. Common file structure")
    add_heading(annex, "3.1 Physical hierarchy")
    add_table_before(doc, annex, ["Element", "Representation"], common["layout"], [4.0, 11.7])
    add_heading(annex, "3.2 Timing rules")
    add_bullets(doc, annex, common["timing_rules"])

    add_title(annex, "4. Common framework records")
    for record in common["common_records"]:
        add_heading(annex, record["name"], 1)
        add_body(annex, "Producer: " + record["producer"], "Producer:")
        add_body(annex, "Condition: " + record["trigger"], "Condition:")
        add_body(annex, "Fields: " + record["fields"], "Fields:")
        add_body(annex, "Interpretation: " + record["meaning"], "Interpretation:")
        add_body(annex, "Status: " + record["stability"], "Status:")

    add_title(annex, "5. Machine-specific components")
    add_heading(annex, "5.1 Selectable inventory")
    add_table_before(doc, annex, ["Id", "Component", "DEBUG output"], machine["components"], [2.0, 4.2, 9.5])
    add_heading(annex, "5.2 Record catalog")
    for record in machine["records"]:
        add_heading(annex, record["name"], 2)
        add_body(annex, "Producer: " + record["producer"], "Producer:")
        add_body(annex, "Fields: " + record["fields"], "Fields:")
        add_body(annex, "Timing: " + record["timing"], "Timing:")
        add_body(annex, "Events: " + record["events"], "Events:")

    add_title(annex, "6. Interpretation method")
    add_heading(annex, "6.1 Recommended procedure")
    add_bullets(doc, annex, common["interpretation_workflow"], numbered=True)
    add_heading(annex, "6.2 Annotated example")
    add_code(annex, machine["example"])
    add_body(annex, machine["example_note"])

    add_title(annex, "7. Limitations and maintenance")
    add_heading(annex, "7.1 Specific limitations")
    add_bullets(doc, annex, machine["limitations"])
    add_heading(annex, "7.2 Maintenance contract")
    add_bullets(doc, annex, common["maintenance"])

    annex_marker = next(p for p in doc.paragraphs if p.text.startswith("[Añada"))
    annex_marker.text = "Appendix A. Provenance and revision"
    annex_marker.style = "Heading 1"
    keep_with_next(annex_marker)
    illustrations = next(p for p in doc.paragraphs if p.text == "Illustrations")
    add_body(illustrations, "Structured sources: .codex/documentation/debug-files/sources/debug-common.yaml and .codex/documentation/debug-files/sources/" + machine["id"] + ".yaml.")
    add_heading(illustrations, "A.1 Watched source files")
    source_rows = [[path, "Common"] for path in common["watched_sources"]] + [[path, machine["id"]] for path in machine["watched_sources"]]
    add_table_before(doc, illustrations, ["Path", "Scope"], source_rows, [12.7, 3.0])
    add_heading(illustrations, "A.2 Revision history")
    add_table_before(doc, illustrations, ["Version", "Year", "Description"], [["1.0", str(common["year"]), "Initial inventory generated and checked against the current C++ producers."]], [2.0, 2.0, 11.7])

    for paragraph in doc.paragraphs:
        if paragraph.style.name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
            keep_with_next(paragraph)
            if paragraph.style.name != "Title":
                remove_page_break_before(paragraph)
    set_update_fields(doc)
    doc.save(output)
    return output


def producer_files() -> list[str]:
    tokens = ("writeCompleteLine", "writeLineData", "writeSimpleLine", "_deepDebugFile", "setDeepDebugFile", "activateDeepDebug")
    result = []
    for base in (ROOT / "src", ROOT / "include"):
        for path in base.rglob("*"):
            if path.suffix.lower() not in (".cpp", ".hpp"):
                continue
            text = path.read_text(encoding="utf-8", errors="ignore")
            if any(token in text for token in tokens):
                result.append(path.relative_to(ROOT).as_posix())
    return sorted(result)


def write_state(common: dict, machines: list[dict]) -> None:
    definitions = [COMMON_PATH, HERE / "generate_debug_docs.py", HERE / "check_debug_docs.py", HERE / "update_debug_doc_fields.ps1"] + [SOURCES / f"{m['id']}.yaml" for m in machines]
    watched = sorted(set(common["watched_sources"] + [p for m in machines for p in m["watched_sources"]]))
    outputs = [ROOT / m["output"] for m in machines]
    state = {
        "schema_version": 1,
        "definitions": {p.relative_to(ROOT).as_posix(): sha256(p) for p in definitions},
        "watched_sources": {p: sha256(ROOT / p) for p in watched},
        "producer_files": producer_files(),
        "outputs": {p.relative_to(ROOT).as_posix(): sha256(p) for p in outputs},
    }
    STATE_PATH.write_text(json.dumps(state, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--machine", choices=[p.stem for p in MACHINE_PATHS], action="append")
    parser.add_argument("--state-only", action="store_true", help="Update generated-state.json without rebuilding DOCX files")
    parser.add_argument("--finalize", action="store_true", help="Mark fields for refresh and update generated-state.json")
    args = parser.parse_args()
    common = load_yaml_json(COMMON_PATH)
    all_machines = [load_yaml_json(path) for path in MACHINE_PATHS]
    if args.finalize:
        for machine in all_machines:
            output = ROOT / machine["output"]
            doc = Document(output)
            normalize_word_field_results(doc)
            clear_update_fields(doc)
            doc.save(output)
        write_state(common, all_machines)
        print(STATE_PATH.relative_to(ROOT).as_posix())
        return 0
    if args.state_only:
        write_state(common, all_machines)
        print(STATE_PATH.relative_to(ROOT).as_posix())
        return 0
    selected = [m for m in all_machines if not args.machine or m["id"] in args.machine]
    for machine in selected:
        output = build_document(common, machine)
        print(output.relative_to(ROOT).as_posix())
    if len(selected) == len(all_machines):
        write_state(common, all_machines)
        print(STATE_PATH.relative_to(ROOT).as_posix())
    else:
        print("State not updated: partial generation must be followed by a full generation.", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
