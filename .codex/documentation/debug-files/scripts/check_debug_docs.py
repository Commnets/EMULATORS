#!/usr/bin/env python3
"""Fail when DEBUG guides are not synchronized with definitions or code."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path
import sys

from docx import Document
from docx.oxml.ns import qn


HERE = Path(__file__).resolve().parent
PACKAGE = HERE.parent
ROOT = PACKAGE.parents[2]
STATE = PACKAGE / "state" / "generated-state.json"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def producer_files() -> list[str]:
    tokens = ("writeCompleteLine", "writeLineData", "writeSimpleLine", "_deepDebugFile", "setDeepDebugFile", "activateDeepDebug")
    result = []
    for base in (ROOT / "src", ROOT / "include"):
        for path in base.rglob("*"):
            if path.suffix.lower() not in (".cpp", ".hpp"):
                continue
            text = path.read_text(encoding="utf-8", errors="ignore")
            if any(token in text for token in tokens):
                result.append(path.relative_to(ROOT).as_posix())
    return sorted(result)


def structural_errors(path: Path) -> list[str]:
    """Enforce the language and style contract requested for every guide."""
    doc = Document(path)
    errors = []
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        ppr = doc.styles[style_name].element.pPr
        if ppr is not None and ppr.find(qn("w:pageBreakBefore")) is not None:
            errors.append(f"{path.name}: {style_name} has pageBreakBefore")
    for paragraph in doc.paragraphs:
        ppr = paragraph._p.pPr
        if paragraph.style.name.startswith("Heading") and ppr is not None and ppr.find(qn("w:pageBreakBefore")) is not None:
            errors.append(f"{path.name}: heading paragraph has pageBreakBefore: {paragraph.text[:60]}")
        if ppr is not None and ppr.find(qn("w:numPr")) is not None and paragraph.style.name != "Normal":
            errors.append(f"{path.name}: numbered paragraph does not use Normal: {paragraph.style.name}")
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    for forbidden in ("Propósito", "Alcance", "Productor:", "Condición:", "Campos:", "Temporización:", "Anexos", "Ilustraciones", "Ecuaciones", "Página", "NO SE ENCUENTRAN"):
        if forbidden in full_text:
            errors.append(f"{path.name}: Spanish document label remains: {forbidden}")
    return errors


def main() -> int:
    if not STATE.exists():
        print(".codex/documentation/debug-files/state/generated-state.json is missing; regenerate the guides.", file=sys.stderr)
        return 1
    state = json.loads(STATE.read_text(encoding="utf-8"))
    errors = []
    for group in ("definitions", "watched_sources", "outputs"):
        for relative, expected in state[group].items():
            path = ROOT / relative
            if not path.exists():
                errors.append(f"Missing {relative}")
            elif sha256(path) != expected:
                errors.append(f"Changed {relative}")
    current_producers = producer_files()
    if current_producers != state["producer_files"]:
        added = sorted(set(current_producers) - set(state["producer_files"]))
        removed = sorted(set(state["producer_files"]) - set(current_producers))
        if added:
            errors.append("New potential producers: " + ", ".join(added))
        if removed:
            errors.append("Removed potential producers: " + ", ".join(removed))
    for relative in state["outputs"]:
        path = ROOT / relative
        if path.exists():
            errors.extend(structural_errors(path))
    if errors:
        print("The DEBUG guides must be reviewed and regenerated:", file=sys.stderr)
        for error in errors:
            print(" - " + error, file=sys.stderr)
        return 1
    print("DEBUG guides synchronized: 6 machine-specific definitions and 6 DOCX files.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
